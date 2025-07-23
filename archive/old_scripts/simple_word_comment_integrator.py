"""
Vereinfachte Word-Kommentar Integration mit python-docx
Nutzt die eingebauten Kommentar-Funktionen
"""

from docx import Document
from docx.shared import RGBColor
from typing import List, Dict, Tuple, Optional
import re
import os


class SimpleWordCommentIntegrator:
    """Vereinfachte Word-Kommentar Integration"""
    
    def __init__(self, document_path: str):
        self.document_path = document_path
        self.document = Document(document_path)
        self.comments_added = 0
    
    def add_track_changes_comments(self, suggestions: List) -> int:
        """Fügt Kommentare mit Track-Changes-ähnlicher Formatierung hinzu"""
        comments_added = 0
        
        for suggestion in suggestions:
            if self._add_track_change_comment(suggestion):
                comments_added += 1
        
        self.comments_added = comments_added
        return comments_added
    
    def _add_track_change_comment(self, suggestion) -> bool:
        """Fügt einen Track-Changes-ähnlichen Kommentar hinzu"""
        try:
            # Finde den passenden Absatz
            target_paragraph = self._find_paragraph_with_text(suggestion.original_text)
            
            if not target_paragraph:
                return False
            
            # Erstelle Kommentar-Text
            comment_text = self._format_track_change_comment(suggestion)
            
            # Methode 1: Füge Kommentar am Ende des Absatzes hinzu
            original_text = target_paragraph.text
            
            # Leere den Absatz
            target_paragraph.clear()
            
            # Füge ursprünglichen Text hinzu
            original_run = target_paragraph.add_run(original_text)
            
            # Füge Kommentar als hervorgehobenen Text hinzu
            comment_run = target_paragraph.add_run(comment_text)
            
            # Formatiere Kommentar wie Word Track Changes
            comment_run.font.color.rgb = RGBColor(197, 90, 17)  # Orange wie Word-Kommentare
            comment_run.font.size = target_paragraph.runs[0].font.size if target_paragraph.runs else None
            
            # Alternativ: Verwende Highlighting
            from docx.enum.text import WD_COLOR_INDEX
            try:
                comment_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            except:
                pass  # Falls nicht unterstützt
            
            return True
            
        except Exception as e:
            print(f"Fehler beim Track-Changes-Kommentar: {e}")
            return False
    
    def add_bubble_comments(self, suggestions: List) -> int:
        """Alternative: Fügt Kommentare als 'Sprechblasen' hinzu"""
        comments_added = 0
        
        for suggestion in suggestions:
            if self._add_bubble_comment(suggestion):
                comments_added += 1
        
        self.comments_added = comments_added
        return comments_added
    
    def _add_bubble_comment(self, suggestion) -> bool:
        """Fügt Kommentar als Sprechblase hinzu"""
        try:
            target_paragraph = self._find_paragraph_with_text(suggestion.original_text)
            
            if not target_paragraph:
                return False
            
            # Erstelle "Sprechblasen"-Kommentar
            bubble_comment = self._format_bubble_comment(suggestion)
            
            # Füge nach dem Absatz einen neuen Kommentar-Absatz ein
            # Finde Index des aktuellen Absatzes
            paragraphs = self.document.paragraphs
            current_index = -1
            
            for i, para in enumerate(paragraphs):
                if para == target_paragraph:
                    current_index = i
                    break
            
            if current_index >= 0:
                # Füge neuen Absatz nach dem aktuellen ein
                new_para = self._insert_paragraph_after(current_index)
                
                # Formatiere als Kommentar-Blase
                comment_run = new_para.add_run(bubble_comment)
                comment_run.font.color.rgb = RGBColor(255, 255, 255)  # Weißer Text
                comment_run.font.size = comment_run.font.size
                comment_run.bold = True
                
                # Setze Absatz-Hintergrund (funktioniert nicht immer)
                try:
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'D73502')  # Orange Hintergrund
                    new_para._element.get_or_add_pPr().append(shading_elm)
                except:
                    # Fallback: Nur Text-Farbe
                    comment_run.font.color.rgb = RGBColor(215, 53, 2)
                
                return True
                
        except Exception as e:
            print(f"Fehler beim Bubble-Kommentar: {e}")
            return False
        
        return False
    
    def _insert_paragraph_after(self, index: int):
        """Fügt einen neuen Absatz nach dem gegebenen Index ein"""
        # Das ist ein Workaround da python-docx keine direkte insert-Funktion hat
        return self.document.add_paragraph()  # Wird am Ende hinzugefügt
    
    def _find_paragraph_with_text(self, search_text: str) -> Optional[object]:
        """Findet Absatz der den Suchtext enthält"""
        search_text = search_text.strip()[:50]  # Erste 50 Zeichen
        
        for paragraph in self.document.paragraphs:
            if search_text.lower() in paragraph.text.lower():
                return paragraph
        
        return None
    
    def _format_track_change_comment(self, suggestion) -> str:
        """Formatiert Kommentar im Track-Changes-Stil"""
        category_map = {
            'grammar': 'GRAMMATIK',
            'style': 'STIL',
            'clarity': 'KLARHEIT', 
            'academic': 'WISSENSCHAFT'
        }
        
        category = category_map.get(suggestion.category.lower(), 'ALLGEMEIN')
        
        comment = f" [KI-{category}: {suggestion.reason} → {suggestion.suggested_text}] "
        return comment
    
    def _format_bubble_comment(self, suggestion) -> str:
        """Formatiert Kommentar als Sprechblase"""
        category_icons = {
            'grammar': '🔴',
            'style': '🟡',
            'clarity': '🟢',
            'academic': '🔵'
        }
        
        icon = category_icons.get(suggestion.category.lower(), '⚪')
        
        comment = f"\n{icon} KI-VERBESSERUNG: {suggestion.suggested_text}\n"
        comment += f"💭 {suggestion.reason}\n"
        
        return comment
    
    def add_margin_comments(self, suggestions: List) -> int:
        """Fügt Kommentare als Rand-Notizen hinzu"""
        comments_added = 0
        
        # Sammle alle Suggestions nach Absätzen
        paragraph_suggestions = {}
        
        for suggestion in suggestions:
            para = self._find_paragraph_with_text(suggestion.original_text)
            if para:
                if para not in paragraph_suggestions:
                    paragraph_suggestions[para] = []
                paragraph_suggestions[para].append(suggestion)
        
        # Füge Kommentare zu den Absätzen hinzu
        for paragraph, para_suggestions in paragraph_suggestions.items():
            if self._add_margin_comment_to_paragraph(paragraph, para_suggestions):
                comments_added += len(para_suggestions)
        
        self.comments_added = comments_added
        return comments_added
    
    def _add_margin_comment_to_paragraph(self, paragraph, suggestions: List) -> bool:
        """Fügt Rand-Kommentare zu einem Absatz hinzu"""
        try:
            # Erstelle zusammengefassten Kommentar für alle Suggestions in diesem Absatz
            margin_comment = "\n📝 KI-VERBESSERUNGSVORSCHLÄGE:\n\n"
            
            for i, suggestion in enumerate(suggestions, 1):
                category_icons = {
                    'grammar': '🔴 GRAMMATIK',
                    'style': '🟡 STIL', 
                    'clarity': '🟢 KLARHEIT',
                    'academic': '🔵 WISSENSCHAFT'
                }
                
                icon = category_icons.get(suggestion.category.lower(), '⚪ ALLGEMEIN')
                
                margin_comment += f"{i}. {icon}\n"
                margin_comment += f"   💡 {suggestion.suggested_text}\n"
                margin_comment += f"   📋 {suggestion.reason}\n\n"
            
            margin_comment += "─" * 40 + "\n"
            
            # Füge Kommentar am Ende des Absatzes hinzu
            original_text = paragraph.text
            
            # Leere und fülle Absatz neu
            paragraph.clear()
            
            # Ursprünglicher Text
            original_run = paragraph.add_run(original_text)
            
            # Rand-Kommentar
            comment_run = paragraph.add_run(margin_comment)
            comment_run.font.color.rgb = RGBColor(128, 128, 128)  # Grau
            comment_run.italic = True
            
            return True
            
        except Exception as e:
            print(f"Fehler beim Rand-Kommentar: {e}")
            return False
    
    def save_document(self, output_path: str) -> bool:
        """Speichert das Dokument"""
        try:
            self.document.save(output_path)
            print(f"✅ Dokument mit Kommentaren gespeichert: {output_path}")
            return True
        except Exception as e:
            print(f"❌ Fehler beim Speichern: {e}")
            return False
    
    def create_backup(self) -> str:
        """Erstellt Backup"""
        backup_path = self.document_path.replace('.docx', '_backup.docx')
        try:
            import shutil
            shutil.copy2(self.document_path, backup_path)
            print(f"🔒 Backup erstellt: {backup_path}")
            return backup_path
        except Exception as e:
            print(f"❌ Fehler beim Backup: {e}")
            return ""


def main():
    """Test alle Kommentar-Stile"""
    from dataclasses import dataclass
    
    @dataclass 
    class MockSuggestion:
        original_text: str
        suggested_text: str
        reason: str
        category: str
        confidence: float
        position: Tuple[int, int]
    
    document_path = '/Users/max/Korrekturtool BA/Volltext_BA_Max Thomsen Kopie.docx'
    
    # Test-Suggestions
    test_suggestions = [
        MockSuggestion(
            original_text="Large Language Models",
            suggested_text="Large Language Models (LLMs)", 
            reason="Abkürzung beim ersten Gebrauch ausschreiben",
            category="academic",
            confidence=0.9,
            position=(100, 120)
        )
    ]
    
    # Teste verschiedene Kommentar-Stile
    styles = [
        ('track_changes', 'add_track_changes_comments'),
        ('bubble', 'add_bubble_comments'),
        ('margin', 'add_margin_comments')
    ]
    
    for style_name, method_name in styles:
        print(f"\n=== Teste {style_name.upper()} Kommentare ===")
        
        integrator = SimpleWordCommentIntegrator(document_path)
        backup_path = integrator.create_backup()
        
        # Rufe entsprechende Methode auf
        method = getattr(integrator, method_name)
        comments_added = method(test_suggestions)
        
        # Speichere mit Style-spezifischem Namen
        output_path = document_path.replace('.docx', f'_{style_name}_comments.docx')
        success = integrator.save_document(output_path)
        
        if success:
            print(f"✅ {comments_added} {style_name} Kommentare hinzugefügt")
        else:
            print(f"❌ Fehler bei {style_name} Kommentaren")


if __name__ == "__main__":
    main()