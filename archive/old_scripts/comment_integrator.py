"""
Kommentar-Integration in Word-Dokumente
Fügt KI-Verbesserungsvorschläge als Inline-Kommentare in das originale Dokument ein
"""

from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement, ns
from typing import List, Dict, Tuple, Optional
import re
import os
from pathlib import Path


class CommentIntegrator:
    """Integriert Kommentare in Word-Dokumente ohne Struktur zu zerstören"""
    
    def __init__(self, document_path: str):
        self.document_path = document_path
        self.document = Document(document_path)
        self.original_text = self._extract_full_text()
        self.paragraph_mapping = self._create_paragraph_mapping()
    
    def _extract_full_text(self) -> str:
        """Extrahiert den vollständigen Text aus dem Dokument"""
        full_text = ""
        for paragraph in self.document.paragraphs:
            full_text += paragraph.text + "\n"
        return full_text
    
    def _create_paragraph_mapping(self) -> Dict[int, Tuple[int, int]]:
        """Erstellt Mapping zwischen Absätzen und Textpositionen"""
        mapping = {}
        current_pos = 0
        
        for i, paragraph in enumerate(self.document.paragraphs):
            text = paragraph.text
            start_pos = current_pos
            end_pos = current_pos + len(text)
            mapping[i] = (start_pos, end_pos)
            current_pos = end_pos + 1  # +1 für Newline
        
        return mapping
    
    def find_paragraph_for_position(self, position: int) -> Optional[int]:
        """Findet den Absatz-Index für eine Textposition"""
        for para_idx, (start, end) in self.paragraph_mapping.items():
            if start <= position <= end:
                return para_idx
        return None
    
    def add_inline_comments(self, suggestions: List) -> int:
        """Fügt Inline-Kommentare basierend auf Suggestions hinzu"""
        comments_added = 0
        
        # Sortiere Suggestions nach Position (rückwärts, um Positionen nicht zu verschieben)
        sorted_suggestions = sorted(suggestions, key=lambda x: x.position[0], reverse=True)
        
        for suggestion in sorted_suggestions:
            if self._add_single_comment(suggestion):
                comments_added += 1
        
        return comments_added
    
    def _add_single_comment(self, suggestion) -> bool:
        """Fügt einen einzelnen Kommentar hinzu"""
        try:
            start_pos, end_pos = suggestion.position
            para_idx = self.find_paragraph_for_position(start_pos)
            
            if para_idx is None:
                return False
            
            paragraph = self.document.paragraphs[para_idx]
            para_start, para_end = self.paragraph_mapping[para_idx]
            
            # Berechne relative Position im Absatz
            relative_start = start_pos - para_start
            relative_end = end_pos - para_start
            
            # Erstelle Kommentar-Text
            comment_text = f" [VORSCHLAG: {suggestion.suggested_text} | {suggestion.reason}] "
            
            # Einfache Methode: Füge Kommentar nach dem zu korrigierenden Text ein
            original_text = paragraph.text
            
            if relative_end <= len(original_text):
                new_text = (
                    original_text[:relative_end] + 
                    comment_text + 
                    original_text[relative_end:]
                )
                
                # Ersetze Paragraph-Text
                paragraph.clear()
                run = paragraph.add_run(new_text)
                
                # Markiere Kommentar-Teil farblich
                self._highlight_comment_in_paragraph(paragraph, comment_text)
                
                return True
        
        except Exception as e:
            print(f"Fehler beim Hinzufügen des Kommentars: {e}")
            
        return False
    
    def _highlight_comment_in_paragraph(self, paragraph, comment_text: str):
        """Hebt Kommentare im Absatz farblich hervor"""
        try:
            # Finde alle Runs im Paragraph
            full_text = paragraph.text
            comment_start = full_text.find(comment_text)
            
            if comment_start >= 0:
                # Teile den Text in Runs auf
                paragraph.clear()
                
                # Text vor Kommentar
                if comment_start > 0:
                    paragraph.add_run(full_text[:comment_start])
                
                # Kommentar (hervorgehoben)
                comment_run = paragraph.add_run(comment_text)
                comment_run.font.color.rgb = RGBColor(255, 0, 0)  # Rot
                comment_run.bold = True
                
                # Text nach Kommentar
                comment_end = comment_start + len(comment_text)
                if comment_end < len(full_text):
                    paragraph.add_run(full_text[comment_end:])
        
        except Exception as e:
            print(f"Fehler beim Hervorheben: {e}")
    
    def add_bracket_comments(self, suggestions: List) -> int:
        """Fügt Kommentare in eckigen Klammern hinzu (einfachere Methode)"""
        comments_added = 0
        
        for suggestion in suggestions:
            try:
                # Finde den entsprechenden Absatz
                start_pos, end_pos = suggestion.position
                para_idx = self.find_paragraph_for_position(start_pos)
                
                if para_idx is not None:
                    paragraph = self.document.paragraphs[para_idx]
                    original_text = paragraph.text
                    
                    # Erstelle Bracket-Kommentar
                    bracket_comment = f" [{suggestion.category.upper()}: {suggestion.reason}]"
                    
                    # Finde die Position des zu kommentierenden Textes
                    para_start = self.paragraph_mapping[para_idx][0]
                    relative_pos = end_pos - para_start
                    
                    if relative_pos <= len(original_text):
                        new_text = (
                            original_text[:relative_pos] + 
                            bracket_comment + 
                            original_text[relative_pos:]
                        )
                        
                        paragraph.clear()
                        paragraph.add_run(new_text)
                        comments_added += 1
            
            except Exception as e:
                print(f"Fehler beim Bracket-Kommentar: {e}")
        
        return comments_added
    
    def save_commented_document(self, output_path: str) -> bool:
        """Speichert das kommentierte Dokument"""
        try:
            self.document.save(output_path)
            print(f"Kommentiertes Dokument gespeichert: {output_path}")
            return True
        except Exception as e:
            print(f"Fehler beim Speichern: {e}")
            return False
    
    def create_backup(self) -> str:
        """Erstellt eine Backup-Kopie des Originaldokuments"""
        backup_path = self.document_path.replace('.docx', '_backup.docx')
        try:
            # Kopiere Original
            import shutil
            shutil.copy2(self.document_path, backup_path)
            print(f"Backup erstellt: {backup_path}")
            return backup_path
        except Exception as e:
            print(f"Fehler beim Backup: {e}")
            return ""
    
    def get_document_stats(self) -> Dict[str, int]:
        """Gibt Statistiken über das Dokument zurück"""
        return {
            'paragraphs': len(self.document.paragraphs),
            'characters': len(self.original_text),
            'words': len(self.original_text.split()),
        }


def main():
    """Test-Funktion"""
    import sys
    from dataclasses import dataclass
    
    # Mock Suggestion für Test
    @dataclass
    class MockSuggestion:
        original_text: str
        suggested_text: str
        reason: str
        category: str
        confidence: float
        position: Tuple[int, int]
    
    if len(sys.argv) != 2:
        print("Usage: python comment_integrator.py <document.docx>")
        return
    
    document_path = sys.argv[1]
    
    if not os.path.exists(document_path):
        print(f"Dokument nicht gefunden: {document_path}")
        return
    
    integrator = CommentIntegrator(document_path)
    stats = integrator.get_document_stats()
    
    print(f"Dokument-Statistiken:")
    print(f"- Absätze: {stats['paragraphs']}")
    print(f"- Wörter: {stats['words']}")
    print(f"- Zeichen: {stats['characters']}")
    
    # Test-Suggestions
    test_suggestions = [
        MockSuggestion(
            original_text="Dies ist ein Testtext",
            suggested_text="Dies ist ein verbesserter Testtext",
            reason="Präzisere Formulierung",
            category="style",
            confidence=0.8,
            position=(0, 21)
        )
    ]
    
    # Backup erstellen
    backup_path = integrator.create_backup()
    
    # Kommentare hinzufügen
    comments_added = integrator.add_bracket_comments(test_suggestions)
    print(f"Kommentare hinzugefügt: {comments_added}")
    
    # Speichern
    output_path = document_path.replace('.docx', '_commented.docx')
    integrator.save_commented_document(output_path)


if __name__ == "__main__":
    main()