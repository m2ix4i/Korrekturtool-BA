"""
Verbesserte Kommentar-Integration mit deutlich sichtbaren Kommentaren
"""

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from typing import List, Dict, Tuple, Optional
import re
import os
from pathlib import Path


class EnhancedCommentIntegrator:
    """Verbesserte Kommentar-Integration mit sichtbaren, farbigen Kommentaren"""
    
    def __init__(self, document_path: str):
        self.document_path = document_path
        self.document = Document(document_path)
        self.original_text = self._extract_full_text()
        self.paragraph_mapping = self._create_paragraph_mapping()
        self.comments_added = 0
    
    def _extract_full_text(self) -> str:
        """Extrahiert den vollständigen Text aus dem Dokument"""
        full_text = ""
        for paragraph in self.document.paragraphs:
            full_text += paragraph.text + "\n"
        return full_text
    
    def _create_paragraph_mapping(self) -> Dict[int, Tuple[int, int]]:
        """Erstellt Mapping zwischen Absätzen und Textpositionen"""
        mapping = {}
        current_pos = 0
        
        for i, paragraph in enumerate(self.document.paragraphs):
            text = paragraph.text
            start_pos = current_pos
            end_pos = current_pos + len(text)
            mapping[i] = (start_pos, end_pos)
            current_pos = end_pos + 1  # +1 für Newline
        
        return mapping
    
    def add_highlighted_comments(self, suggestions: List) -> int:
        """Fügt farbig hervorgehobene Kommentare hinzu"""
        comments_added = 0
        
        # Sortiere Suggestions nach Position (rückwärts)
        sorted_suggestions = sorted(suggestions, key=lambda x: x.position[0], reverse=True)
        
        for suggestion in sorted_suggestions:
            if self._add_highlighted_comment(suggestion):
                comments_added += 1
        
        self.comments_added = comments_added
        return comments_added
    
    def _add_highlighted_comment(self, suggestion) -> bool:
        """Fügt einen farbig hervorgehobenen Kommentar hinzu"""
        try:
            # Finde den besten Absatz für den Kommentar
            para_idx = self._find_best_paragraph_for_suggestion(suggestion)
            
            if para_idx is None:
                return False
            
            paragraph = self.document.paragraphs[para_idx]
            original_text = paragraph.text
            
            # Erstelle deutlichen Kommentar mit Kategorie-Färbung
            category_colors = {
                'grammar': '🔴',
                'style': '🟡', 
                'clarity': '🟢',
                'academic': '🔵'
            }
            
            icon = category_colors.get(suggestion.category.lower(), '⚪')
            
            comment_text = f" {icon}[{suggestion.category.upper()}: {suggestion.reason}]{icon} "
            
            # Füge Kommentar am Ende des Absatzes hinzu (sichtbarer)
            new_text = original_text + comment_text
            
            # Ersetze Paragraph-Inhalt
            paragraph.clear()
            
            # Füge ursprünglichen Text hinzu
            original_run = paragraph.add_run(original_text)
            
            # Füge Kommentar hervorgehoben hinzu
            comment_run = paragraph.add_run(comment_text)
            comment_run.font.color.rgb = RGBColor(255, 0, 0)  # Rot
            comment_run.bold = True
            comment_run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # Gelb hervorheben
            
            return True
            
        except Exception as e:
            print(f"Fehler beim Hinzufügen des Kommentars: {e}")
            return False
    
    def _find_best_paragraph_for_suggestion(self, suggestion) -> Optional[int]:
        """Findet den besten Absatz für eine Suggestion"""
        original_text = suggestion.original_text[:30]  # Erste 30 Zeichen
        
        # Suche in allen Absätzen
        for i, paragraph in enumerate(self.document.paragraphs):
            if original_text.lower() in paragraph.text.lower():
                return i
        
        # Fallback: Verwende die ursprüngliche Positionsmethode
        start_pos, _ = suggestion.position
        for para_idx, (para_start, para_end) in self.paragraph_mapping.items():
            if para_start <= start_pos <= para_end:
                return para_idx
                
        return None
    
    def add_summary_at_end(self, total_suggestions: int):
        """Fügt eine Zusammenfassung der Kommentare am Ende hinzu"""
        try:
            # Füge neue Seite hinzu
            self.document.add_page_break()
            
            # Überschrift
            heading = self.document.add_heading('🎓 KI-Korrekturbericht', 1)
            
            # Zusammenfassung
            summary = self.document.add_paragraph()
            summary.add_run(f"Dieses Dokument wurde mit dem Bachelorarbeit-Korrekturtool analysiert.\n\n")
            summary.add_run(f"📊 Analyseergebnisse:\n").bold = True
            summary.add_run(f"• Gefundene Verbesserungsvorschläge: {total_suggestions}\n")
            summary.add_run(f"• Erfolgreich integrierte Kommentare: {self.comments_added}\n")
            summary.add_run(f"• Verwendete KI: Google Gemini-1.5-flash\n\n")
            
            # Legende
            legend = self.document.add_paragraph()
            legend.add_run("📋 Kommentar-Kategorien:\n").bold = True
            legend.add_run("🔴 GRAMMAR: Grammatikalische Korrekturen\n")
            legend.add_run("🟡 STYLE: Stilistische Verbesserungen\n") 
            legend.add_run("🟢 CLARITY: Klarheit und Verständlichkeit\n")
            legend.add_run("🔵 ACADEMIC: Wissenschaftlicher Ausdruck\n\n")
            
            # Hinweis
            note = self.document.add_paragraph()
            note.add_run("💡 Hinweis: ").bold = True
            note.add_run("Alle Kommentare sind Verbesserungsvorschläge. Prüfen Sie jeden Vorschlag vor der Übernahme.")
            
        except Exception as e:
            print(f"Fehler beim Hinzufügen der Zusammenfassung: {e}")
    
    def save_document(self, output_path: str) -> bool:
        """Speichert das kommentierte Dokument"""
        try:
            self.document.save(output_path)
            print(f"✅ Verbessertes Dokument gespeichert: {output_path}")
            return True
        except Exception as e:
            print(f"❌ Fehler beim Speichern: {e}")
            return False
    
    def create_backup(self) -> str:
        """Erstellt eine Backup-Kopie des Originaldokuments"""
        backup_path = self.document_path.replace('.docx', '_backup.docx')
        try:
            import shutil
            shutil.copy2(self.document_path, backup_path)
            print(f"🔒 Backup erstellt: {backup_path}")
            return backup_path
        except Exception as e:
            print(f"❌ Fehler beim Backup: {e}")
            return ""


def main():
    """Test-Funktion"""
    from dataclasses import dataclass
    
    @dataclass 
    class MockSuggestion:
        original_text: str
        suggested_text: str
        reason: str
        category: str
        confidence: float
        position: Tuple[int, int]
    
    document_path = '/Users/max/Korrekturtool BA/Volltext_BA_Max Thomsen Kopie.docx'
    
    integrator = EnhancedCommentIntegrator(document_path)
    
    # Test-Suggestions
    test_suggestions = [
        MockSuggestion(
            original_text="Large Language Models",
            suggested_text="Large Language Models (LLMs)", 
            reason="Abkürzung beim ersten Gebrauch ausschreiben",
            category="academic",
            confidence=0.9,
            position=(100, 120)
        )
    ]
    
    # Backup erstellen
    backup_path = integrator.create_backup()
    
    # Kommentare hinzufügen  
    comments_added = integrator.add_highlighted_comments(test_suggestions)
    
    # Zusammenfassung hinzufügen
    integrator.add_summary_at_end(len(test_suggestions))
    
    # Speichern
    output_path = document_path.replace('.docx', '_enhanced.docx')
    success = integrator.save_document(output_path)
    
    if success:
        print(f"🎉 Test erfolgreich: {comments_added} Kommentare hinzugefügt")


if __name__ == "__main__":
    main()