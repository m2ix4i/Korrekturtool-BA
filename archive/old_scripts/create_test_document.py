#!/usr/bin/env python3
"""
Erstellt ein Test-Word-Dokument für den Parser
"""

from docx import Document
from docx.shared import Inches

def create_test_document():
    doc = Document()
    
    # Titel
    title = doc.add_heading('Testdokument für Korrekturtool', 0)
    
    # Einleitung
    doc.add_heading('1. Einleitung', level=1)
    doc.add_paragraph(
        'Dies ist ein Testdokument für das Bachelorarbeit-Korrekturtool. '
        'Es enthält verschiedene Textelemente wie Überschriften, Absätze, '
        'und andere strukturelle Komponenten die in einer typischen '
        'wissenschaftlichen Arbeit vorkommen.'
    )
    
    # Hauptteil
    doc.add_heading('2. Hauptteil', level=1)
    doc.add_paragraph(
        'Dieser Abschnitt enthält den Hauptinhalt der Arbeit. Hier können '
        'verschiedene grammatikalische und stilistische Fehler auftreten, '
        'die vom KI-System erkannt und korrigiert werden sollen. Zum Beispiel '
        'könnte ein Satz zu lang sein oder unklare Formulierungen enthalten.'
    )
    
    doc.add_paragraph(
        'Ein weiterer Absatz mit potentiellen Verbesserungsmöglichkeiten. '
        'Die Satzstruktur könnte optimiert werden und der wissenschaftliche '
        'Ausdruck könnte präziser formuliert sein.'
    )
    
    # Unterüberschrift
    doc.add_heading('2.1 Methodologie', level=2)
    doc.add_paragraph(
        'Die angewandte Methodik basiert auf etablierten Verfahren der '
        'Textanalyse. Hierbei werden moderne KI-Technologien eingesetzt '
        'um eine umfassende Bewertung des Textes zu ermöglichen.'
    )
    
    # Abbildung (Platzhalter)
    doc.add_paragraph().add_run().add_picture = lambda: None  # Würde normalerweise Bild einfügen
    caption = doc.add_paragraph('Abbildung 1: Beispielabbildung für Testzwecke')
    caption.style = doc.styles['Caption']
    
    # Fazit
    doc.add_heading('3. Fazit', level=1)
    doc.add_paragraph(
        'Zusammenfassend lässt sich sagen, dass dieses Testdokument '
        'verschiedene Herausforderungen für das Korrektursystem darstellt. '
        'Die erfolgreiche Verarbeitung wird zeigen, ob das System '
        'korrekt funktioniert.'
    )
    
    # Speichern
    doc.save('/Users/max/Korrekturtool BA/data/test_document.docx')
    print("Test-Dokument erstellt: data/test_document.docx")

if __name__ == "__main__":
    create_test_document()