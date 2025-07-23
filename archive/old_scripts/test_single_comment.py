#!/usr/bin/env python3
"""
Test für einen einzelnen Word-Kommentar
Debuggt die XML-Struktur und Kommentar-Integration Schritt für Schritt
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET
import datetime
import shutil


def create_single_test_comment():
    """Erstellt genau einen Test-Kommentar mit definiertem Inhalt"""
    
    # Lade Originaldokument
    source_path = '/Users/max/Korrekturtool BA/Volltext_BA_Max Thomsen Kopie.docx'
    test_path = '/Users/max/Korrekturtool BA/TEST_SINGLE_COMMENT.docx'
    
    # Kopiere für Test
    shutil.copy2(source_path, test_path)
    
    doc = Document(test_path)
    
    print("=== SINGLE COMMENT TEST ===")
    print(f"Originaldokument kopiert: {test_path}")
    
    # Definiere Test-Kommentar-Daten
    comment_id = "1"
    comment_author = "Max Test"
    comment_initials = "MT"
    comment_date = "2025-01-22T14:00:00Z"
    comment_text = "Dies ist ein Test-Kommentar mit definiertem Inhalt."
    
    print(f"Kommentar-ID: {comment_id}")
    print(f"Author: {comment_author}")
    print(f"Text: {comment_text}")
    
    # 1. Füge Comment-Range zum ersten Paragraph hinzu
    first_paragraph = doc.paragraphs[10]  # Nehme Paragraph 10 für Test
    print(f"Ziel-Paragraph: {first_paragraph.text[:50]}...")
    
    # Füge Comment-Range-Elemente hinzu
    p_elem = first_paragraph._element
    
    # CommentRangeStart
    comment_start = OxmlElement('w:commentRangeStart')
    comment_start.set(qn('w:id'), comment_id)
    p_elem.insert(0, comment_start)
    print("✓ CommentRangeStart hinzugefügt")
    
    # CommentRangeEnd
    comment_end = OxmlElement('w:commentRangeEnd')
    comment_end.set(qn('w:id'), comment_id)
    p_elem.append(comment_end)
    print("✓ CommentRangeEnd hinzugefügt")
    
    # CommentReference
    comment_ref_run = OxmlElement('w:r')
    comment_reference = OxmlElement('w:commentReference')
    comment_reference.set(qn('w:id'), comment_id)
    comment_ref_run.append(comment_reference)
    p_elem.append(comment_ref_run)
    print("✓ CommentReference hinzugefügt")
    
    # 2. Erstelle comments.xml manuell (sehr einfach)
    comments_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:comment w:id="{comment_id}" w:author="{comment_author}" w:date="{comment_date}" w:initials="{comment_initials}">
        <w:p>
            <w:r>
                <w:t>{comment_text}</w:t>
            </w:r>
        </w:p>
    </w:comment>
</w:comments>'''
    
    print("✓ Comments-XML generiert")
    
    # 3. Füge comments.xml zur Document-Struktur hinzu
    try:
        package = doc.part.package
        
        # Erstelle comments.xml Part
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        
        comments_uri = PackURI('/word/comments.xml')
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
        
        # Erstelle Part
        comments_part = Part(
            partname=comments_uri,
            content_type=content_type,
            blob=comments_xml.encode('utf-8'),
            package=package
        )
        
        # Füge zum Package hinzu
        package._parts[comments_uri] = comments_part
        print("✓ Comments-Part zum Package hinzugefügt")
        
        # Erstelle Relationship
        rel_type = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
        doc.part.relate_to(comments_part, rel_type)
        print("✓ Relationship erstellt")
        
    except Exception as e:
        print(f"❌ Fehler bei comments.xml: {e}")
        import traceback
        traceback.print_exc()
        return False
    
    # 4. Speichere Dokument
    try:
        doc.save(test_path)
        print(f"✅ Test-Dokument gespeichert: {test_path}")
        print("\n💡 Öffnen Sie das Dokument in Word und prüfen Sie:")
        print("   1. Gibt es einen XML-Fehler?")
        print("   2. Ist der Kommentar sichtbar? (Überprüfen > Kommentare)")
        print("   3. Hat der Kommentar den richtigen Inhalt?")
        print("   4. Ist der Author korrekt?")
        return True
        
    except Exception as e:
        print(f"❌ Fehler beim Speichern: {e}")
        import traceback
        traceback.print_exc()
        return False


def debug_existing_comments():
    """Debuggt ein existierendes Dokument mit Kommentaren"""
    
    test_path = '/Users/max/Korrekturtool BA/TEST_SINGLE_COMMENT.docx'
    
    try:
        doc = Document(test_path)
        package = doc.part.package
        
        print("=== DEBUG EXISTING COMMENTS ===")
        
        # Prüfe alle Parts
        print("Package Parts:")
        for uri, part in package._parts.items():
            print(f"  {uri}: {part.content_type}")
            
        # Suche nach comments.xml
        comments_part = None
        for uri, part in package._parts.items():
            if 'comments' in str(uri):
                comments_part = part
                print(f"✓ Comments-Part gefunden: {uri}")
                break
        
        if comments_part:
            # Zeige comments.xml Inhalt
            comments_xml = comments_part.blob.decode('utf-8')
            print("Comments-XML Inhalt:")
            print(comments_xml)
        else:
            print("❌ Keine comments.xml gefunden")
            
        # Prüfe Relationships
        print("Document Relationships:")
        for rel_id, rel in doc.part.rels.items():
            print(f"  {rel_id}: {rel.reltype} -> {rel.target_ref}")
            
    except Exception as e:
        print(f"❌ Debug-Fehler: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    success = create_single_test_comment()
    
    if success:
        print("\n" + "="*50)
        debug_existing_comments()