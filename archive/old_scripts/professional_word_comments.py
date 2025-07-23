"""
Professionelle Word-Kommentar Integration
Erstellt echte Word-Kommentare durch direkte XML-Manipulation
"""

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import RGBColor
from typing import List, Dict, Tuple, Optional
import xml.etree.ElementTree as ET
import re
import os
import uuid
import datetime


class ProfessionalWordCommentIntegrator:
    """Erstellt professionelle Word-Kommentare mit XML-Manipulation"""
    
    def __init__(self, document_path: str):
        self.document_path = document_path
        self.document = Document(document_path)
        self.comment_id_counter = 1
        self.comments = []  # Store comments data
        
    def add_professional_comments(self, suggestions: List) -> int:
        """Fügt professionelle Word-Kommentare hinzu"""
        comments_added = 0
        
        for suggestion in suggestions:
            if self._add_professional_comment(suggestion):
                comments_added += 1
                
        # Nach dem Hinzufügen aller Kommentare, erstelle die Comments-Part
        if comments_added > 0:
            self._create_comments_part()
        
        return comments_added
    
    def _add_professional_comment(self, suggestion) -> bool:
        """Fügt einen professionellen Kommentar hinzu"""
        try:
            # Finde Text-Position
            target_paragraph, text_position = self._find_text_in_document(suggestion.original_text)
            
            if not target_paragraph:
                return False
            
            comment_id = str(self.comment_id_counter)
            self.comment_id_counter += 1
            
            # Speichere Kommentar-Daten
            comment_data = {
                'id': comment_id,
                'text': self._format_professional_comment(suggestion),
                'author': 'KI-Korrekturtool',
                'date': datetime.datetime.now().isoformat(),
                'category': suggestion.category
            }
            self.comments.append(comment_data)
            
            # Füge Kommentar-Markierung zum Paragraph hinzu
            self._add_comment_markup_to_paragraph(target_paragraph, comment_id, suggestion.original_text)
            
            return True
            
        except Exception as e:
            print(f"Fehler beim professionellen Kommentar: {e}")
            return False
    
    def _find_text_in_document(self, search_text: str) -> Tuple[Optional[object], Optional[int]]:
        """Findet Text im Dokument"""
        search_text = search_text.strip()[:30]  # Erste 30 Zeichen für Suche
        
        for paragraph in self.document.paragraphs:
            para_text = paragraph.text.lower()
            search_lower = search_text.lower()
            
            if search_lower in para_text:
                position = para_text.find(search_lower)
                return paragraph, position
                
        return None, None
    
    def _format_professional_comment(self, suggestion) -> str:
        """Formatiert professionellen Kommentar"""
        category_names = {
            'grammar': 'Grammatik',
            'style': 'Stil',
            'clarity': 'Klarheit', 
            'academic': 'Wissenschaftlicher Ausdruck'
        }
        
        category = category_names.get(suggestion.category.lower(), 'Allgemein')
        
        comment = f"KI-Verbesserung ({category})\n\n"
        comment += f"Vorschlag: {suggestion.suggested_text}\n\n"
        comment += f"Begründung: {suggestion.reason}\n\n"
        comment += f"Vertrauen: {suggestion.confidence:.0%}"
        
        return comment
    
    def _add_comment_markup_to_paragraph(self, paragraph, comment_id: str, original_text: str):
        """Fügt Kommentar-Markup zum Paragraph hinzu"""
        try:
            # Vereinfachter Ansatz: Markiere den gesamten Paragraph
            p_elem = paragraph._element
            
            # Füge commentRangeStart vor dem ersten Run hinzu
            comment_start = OxmlElement('w:commentRangeStart')
            comment_start.set(qn('w:id'), comment_id)
            
            # Füge am Anfang des Paragraphs ein
            if len(p_elem) > 0:
                p_elem.insert(0, comment_start)
            else:
                p_elem.append(comment_start)
            
            # Füge commentRangeEnd nach dem letzten Run hinzu
            comment_end = OxmlElement('w:commentRangeEnd')
            comment_end.set(qn('w:id'), comment_id)
            p_elem.append(comment_end)
            
            # Füge commentReference hinzu
            comment_ref_run = OxmlElement('w:r')
            comment_ref = OxmlElement('w:commentReference') 
            comment_ref.set(qn('w:id'), comment_id)
            comment_ref_run.append(comment_ref)
            p_elem.append(comment_ref_run)
            
        except Exception as e:
            print(f"Fehler beim Comment-Markup: {e}")
    
    def _create_comments_part(self):
        """Erstellt die Comments-Part im Word-Dokument"""
        try:
            # Erstelle Comments-XML
            comments_xml = self._build_comments_xml()
            
            # Füge zur Document-Struktur hinzu
            package = self.document.part.package
            
            # Erstelle comments.xml Part
            comments_part = package.create_part(
                '/word/comments.xml',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
                comments_xml.encode('utf-8')
            )
            
            # Erstelle Relationship
            self.document.part.relate_to(
                comments_part,
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
            )
            
            print(f"✅ Comments-Part mit {len(self.comments)} Kommentaren erstellt")
            
        except Exception as e:
            print(f"Fehler beim Erstellen der Comments-Part: {e}")
    
    def _build_comments_xml(self) -> str:
        """Erstellt das Comments-XML"""
        xml_parts = []
        xml_parts.append('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>')
        xml_parts.append(f'<w:comments {nsdecls("w")}>')
        
        for comment in self.comments:
            # Farbe nach Kategorie
            color_map = {
                'grammar': 'DC143C',    # Crimson
                'style': 'FF8C00',      # Dark Orange
                'clarity': '32CD32',    # Lime Green
                'academic': '4169E1'    # Royal Blue
            }
            color = color_map.get(comment['category'], '808080')  # Default gray
            
            xml_parts.append(f'''
    <w:comment w:id="{comment['id']}" w:author="{comment['author']}" w:date="{comment['date']}" w:initials="KI">
        <w:p>
            <w:pPr>
                <w:pStyle w:val="CommentText"/>
            </w:pPr>
            <w:r>
                <w:rPr>
                    <w:color w:val="{color}"/>
                    <w:sz w:val="18"/>
                </w:rPr>
                <w:t>{self._escape_xml(comment['text'])}</w:t>
            </w:r>
        </w:p>
    </w:comment>''')
        
        xml_parts.append('</w:comments>')
        
        return ''.join(xml_parts)
    
    def _escape_xml(self, text: str) -> str:
        """Escaped XML-Sonderzeichen"""
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;').replace("'", '&apos;')
    
    def add_review_comments(self, suggestions: List) -> int:
        """Alternative: Fügt Review-ähnliche Kommentare hinzu"""
        comments_added = 0
        
        for suggestion in suggestions:
            if self._add_review_comment(suggestion):
                comments_added += 1
        
        return comments_added
    
    def _add_review_comment(self, suggestion) -> bool:
        """Fügt Review-Kommentar hinzu (sichtbarer Ansatz)"""
        try:
            target_paragraph = None
            
            # Finde Paragraph mit Text
            for paragraph in self.document.paragraphs:
                if suggestion.original_text[:20].lower() in paragraph.text.lower():
                    target_paragraph = paragraph
                    break
            
            if not target_paragraph:
                return False
            
            # Erstelle Review-Kommentar am Ende des Absatzes
            review_comment = self._format_review_comment(suggestion)
            
            # Ursprünglichen Text beibehalten und Kommentar anhängen
            original_text = target_paragraph.text
            
            # Leere Paragraph
            target_paragraph.clear()
            
            # Füge ursprünglichen Text hinzu
            original_run = target_paragraph.add_run(original_text)
            
            # Füge Review-Kommentar hinzu
            comment_run = target_paragraph.add_run(review_comment)
            
            # Formatiere als Word-Review-Kommentar
            comment_run.font.color.rgb = RGBColor(220, 20, 60)  # Crimson
            comment_run.bold = True
            
            # Versuche Highlighting (falls unterstützt)
            try:
                from docx.enum.text import WD_COLOR_INDEX
                comment_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            except:
                pass
            
            return True
            
        except Exception as e:
            print(f"Fehler beim Review-Kommentar: {e}")
            return False
    
    def _format_review_comment(self, suggestion) -> str:
        """Formatiert Review-Kommentar"""
        category_icons = {
            'grammar': '📝',
            'style': '✨', 
            'clarity': '💡',
            'academic': '🎓'
        }
        
        icon = category_icons.get(suggestion.category.lower(), '📋')
        
        comment = f" {icon}[REVIEW: {suggestion.suggested_text} - {suggestion.reason}]"
        return comment
    
    def save_document(self, output_path: str) -> bool:
        """Speichert das Dokument"""
        try:
            self.document.save(output_path)
            print(f"✅ Professionelles Dokument gespeichert: {output_path}")
            return True
        except Exception as e:
            print(f"❌ Fehler beim Speichern: {e}")
            return False
    
    def create_backup(self) -> str:
        """Erstellt Backup"""
        backup_path = self.document_path.replace('.docx', '_backup.docx')
        try:
            import shutil
            shutil.copy2(self.document_path, backup_path)
            print(f"🔒 Backup erstellt: {backup_path}")
            return backup_path
        except Exception as e:
            print(f"❌ Backup-Fehler: {e}")
            return ""


def main():
    """Test professionelle Kommentare"""
    from dataclasses import dataclass
    
    @dataclass 
    class MockSuggestion:
        original_text: str
        suggested_text: str
        reason: str
        category: str
        confidence: float
        position: Tuple[int, int]
    
    document_path = '/Users/max/Korrekturtool BA/Volltext_BA_Max Thomsen Kopie.docx'
    
    # Test-Suggestions
    test_suggestions = [
        MockSuggestion(
            original_text="Large Language Models",
            suggested_text="Large Language Models (LLMs)", 
            reason="Abkürzung beim ersten Gebrauch ausschreiben",
            category="academic",
            confidence=0.9,
            position=(100, 120)
        ),
        MockSuggestion(
            original_text="Relevanz von KI",
            suggested_text="Relevanz künstlicher Intelligenz", 
            reason="Abkürzung ausschreiben für bessere Verständlichkeit",
            category="clarity",
            confidence=0.8,
            position=(200, 220)
        )
    ]
    
    print("=== Teste PROFESSIONELLE Word-Kommentare ===")
    
    integrator = ProfessionalWordCommentIntegrator(document_path)
    backup_path = integrator.create_backup()
    
    # Teste professionelle Kommentare
    professional_comments = integrator.add_professional_comments(test_suggestions)
    
    # Speichere professionelle Version
    prof_output = document_path.replace('.docx', '_PROFESSIONAL_comments.docx')
    success1 = integrator.save_document(prof_output)
    
    print(f"✅ {professional_comments} professionelle Kommentare hinzugefügt")
    
    print("\n=== Teste REVIEW-Style Kommentare ===")
    
    # Teste Review-Style (einfacher, aber sichtbar)
    integrator2 = ProfessionalWordCommentIntegrator(document_path)
    review_comments = integrator2.add_review_comments(test_suggestions)
    
    review_output = document_path.replace('.docx', '_REVIEW_comments.docx')
    success2 = integrator2.save_document(review_output)
    
    print(f"✅ {review_comments} Review-Kommentare hinzugefügt")
    
    if success1 and success2:
        print(f"\n🎉 Beide Kommentar-Stile erfolgreich getestet!")
        print(f"   Professional: {prof_output}")
        print(f"   Review: {review_output}")


if __name__ == "__main__":
    main()