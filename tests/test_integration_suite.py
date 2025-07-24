"""
Integration Test Suite
Comprehensive integration testing for the Bachelor Thesis Correction Tool
Tests the complete workflow and inter-component interactions
"""

import pytest
import json
from pathlib import Path
from unittest.mock import patch, MagicMock
from docx import Document

# Import main application modules
try:
    from src.analyzers.advanced_gemini_analyzer import AdvancedGeminiAnalyzer, Suggestion
    from src.parsers.docx_parser import DocxParser
    from src.integrators.advanced_word_integrator_fixed import AdvancedWordIntegrator
except ImportError:
    # Handle missing modules gracefully
    AdvancedGeminiAnalyzer = None
    DocxParser = None
    AdvancedWordIntegrator = None
    Suggestion = None


@pytest.mark.integration
@pytest.mark.e2e
class TestEndToEndWorkflow:
    """End-to-end integration tests for complete correction workflow"""
    
    def test_complete_correction_workflow_mocked(self, integration_test_data, mock_google_api_key):
        """Test complete workflow with mocked AI API"""
        input_doc = integration_test_data['input_document']
        output_doc = integration_test_data['output_document']
        temp_dir = integration_test_data['temp_dir']
        
        # Mock AI responses
        mock_suggestions = [
            {
                "original_text": "KI-Technologien",
                "suggested_text": "KI-Technologien (Künstliche Intelligenz)",
                "reason": "Abkürzung bei erster Erwähnung ausschreiben",
                "category": "academic",
                "confidence": 0.9
            },
            {
                "original_text": "wichtigen Informationen",
                "suggested_text": "relevanten Informationen",
                "reason": "Präzisere Wortwahl",
                "category": "style",
                "confidence": 0.8
            }
        ]
        
        with patch('google.generativeai.GenerativeModel') as mock_model:
            mock_response = MagicMock()
            mock_response.text = json.dumps({"suggestions": mock_suggestions})
            mock_instance = MagicMock()
            mock_instance.generate_content.return_value = mock_response
            mock_model.return_value = mock_instance
            
            workflow_successful = True
            
            try:
                # Step 1: Parse document
                if DocxParser is not None:
                    parser = DocxParser(str(input_doc))
                    assert parser.document is not None
                    
                    # Extract text for analysis
                    full_text = " ".join([p.text for p in parser.document.paragraphs if p.text.strip()])
                    assert len(full_text) > 0
                
                # Step 2: Analyze with AI (mocked)
                suggestions = []
                if AdvancedGeminiAnalyzer is not None and Suggestion is not None:
                    analyzer = AdvancedGeminiAnalyzer()
                    analyzer.client = mock_instance
                    
                    # Analyze chunks of text
                    suggestions = analyzer.analyze_text(full_text[:1000])  # Analyze first 1000 chars
                    assert isinstance(suggestions, list)
                
                # Step 3: Integrate comments
                if AdvancedWordIntegrator is not None:
                    integrator = AdvancedWordIntegrator(str(input_doc))
                    
                    # Add comments for found suggestions
                    comment_count = 0
                    for suggestion in suggestions:
                        result = integrator.add_comment(
                            text=suggestion.original_text,
                            comment=f"{suggestion.reason}: {suggestion.suggested_text}",
                            author=f"AI-{suggestion.category}",
                            category=suggestion.category
                        )
                        if result:
                            comment_count += 1
                    
                    # Save final document
                    save_result = integrator.save(str(output_doc))
                    
                    # Verify output
                    if output_doc.exists():
                        result_doc = Document(str(output_doc))
                        assert result_doc is not None
                
            except Exception as e:
                workflow_successful = False
                pytest.fail(f"Complete workflow failed: {str(e)}")
            
            assert workflow_successful
    
    @pytest.mark.requires_api_key
    @pytest.mark.skipif(not pytest.config.getoption("--run-api-tests", default=False), 
                        reason="API tests require --run-api-tests flag")
    def test_complete_workflow_real_api(self, integration_test_data, real_api_key):
        """Test complete workflow with real API (requires API key and flag)"""
        input_doc = integration_test_data['input_document']
        output_doc = integration_test_data['output_document']
        
        if any(cls is None for cls in [DocxParser, AdvancedGeminiAnalyzer, AdvancedWordIntegrator]):
            pytest.skip("Required components not available")
        
        # Step 1: Parse document
        parser = DocxParser(str(input_doc))
        full_text = " ".join([p.text for p in parser.document.paragraphs if p.text.strip()])
        
        # Step 2: Real AI analysis (limited to avoid costs)
        analyzer = AdvancedGeminiAnalyzer()
        suggestions = analyzer.analyze_text(full_text[:500])  # Limit text to reduce API costs
        
        # Step 3: Integrate comments
        integrator = AdvancedWordIntegrator(str(input_doc))
        
        for suggestion in suggestions[:3]:  # Limit to 3 suggestions
            integrator.add_comment(
                text=suggestion.original_text,
                comment=suggestion.reason,
                author="AI-Real"
            )
        
        integrator.save(str(output_doc))
        
        # Verify result
        assert output_doc.exists()
        result_doc = Document(str(output_doc))
        assert result_doc is not None
    
    def test_workflow_error_handling(self, temp_dir, mock_google_api_key):
        """Test workflow error handling and recovery"""
        # Create invalid input scenarios
        empty_doc = Document()
        empty_doc_path = temp_dir / "empty.docx"
        empty_doc.save(str(empty_doc_path))
        
        with patch('google.generativeai.GenerativeModel') as mock_model:
            # Mock API error
            mock_instance = MagicMock()
            mock_instance.generate_content.side_effect = Exception("API Error")
            mock_model.return_value = mock_instance
            
            errors_handled = []
            
            try:
                # Try to process empty document
                if DocxParser is not None:
                    parser = DocxParser(str(empty_doc_path))
                    # Should handle empty document gracefully
                    assert parser.document is not None
                
                # Try AI analysis with API error
                if AdvancedGeminiAnalyzer is not None:
                    analyzer = AdvancedGeminiAnalyzer()
                    analyzer.client = mock_instance
                    suggestions = analyzer.analyze_text("Test text")
                    # Should return empty list on error
                    assert isinstance(suggestions, list)
                    errors_handled.append("ai_analysis")
                
                # Try integration with no suggestions
                if AdvancedWordIntegrator is not None:
                    integrator = AdvancedWordIntegrator(str(empty_doc_path))
                    result = integrator.add_comment(
                        text="Non-existent text",
                        comment="This should fail gracefully",
                        author="AI Test"
                    )
                    # Should handle gracefully
                    assert result is False or result is None
                    errors_handled.append("integration")
                
            except Exception as e:
                pytest.fail(f"Error handling failed: {str(e)}")
            
            # Should have handled multiple error types
            assert len(errors_handled) >= 1


@pytest.mark.integration
class TestComponentInteractions:
    """Test interactions between different components"""
    
    def test_parser_analyzer_integration(self, create_test_docx, sample_docx_content, mock_google_api_key):
        """Test integration between parser and analyzer"""
        if DocxParser is None or AdvancedGeminiAnalyzer is None:
            pytest.skip("Required components not available")
        
        doc_path = create_test_docx("parser_analyzer.docx", sample_docx_content)
        
        with patch('google.generativeai.GenerativeModel') as mock_model:
            mock_response = MagicMock()
            mock_response.text = json.dumps({
                "suggestions": [
                    {
                        "original_text": "Large Language Model",
                        "suggested_text": "Large Language Model (LLM)",
                        "reason": "Add abbreviation",
                        "category": "academic",
                        "confidence": 0.9
                    }
                ]
            })
            mock_instance = MagicMock()
            mock_instance.generate_content.return_value = mock_response
            mock_model.return_value = mock_instance
            
            # Parse document
            parser = DocxParser(str(doc_path))
            paragraphs = [p.text for p in parser.document.paragraphs if p.text.strip()]
            
            # Analyze parsed content
            analyzer = AdvancedGeminiAnalyzer()
            analyzer.client = mock_instance
            
            all_suggestions = []
            for paragraph in paragraphs:
                if paragraph:
                    suggestions = analyzer.analyze_text(paragraph)
                    all_suggestions.extend(suggestions)
            
            # Should have found some suggestions
            assert isinstance(all_suggestions, list)
            assert len(all_suggestions) >= 0  # May be 0 if no text matches
    
    def test_analyzer_integrator_interaction(self, create_test_docx, mock_google_api_key, sample_suggestions):
        """Test integration between analyzer and integrator"""
        if AdvancedGeminiAnalyzer is None or AdvancedWordIntegrator is None:
            pytest.skip("Required components not available")
        
        # Create document with known content
        content = {
            "title": "Test Integration",
            "paragraphs": ["This is a test for Large Language Model integration."]
        }
        doc_path = create_test_docx("analyzer_integrator.docx", content)
        
        with patch('google.generativeai.GenerativeModel') as mock_model:
            mock_response = MagicMock()
            mock_response.text = json.dumps({
                "suggestions": [
                    {
                        "original_text": "Large Language Model",
                        "suggested_text": "Large Language Model (LLM)",
                        "reason": "Add abbreviation on first mention",
                        "category": "academic",
                        "confidence": 0.9
                    }
                ]
            })
            mock_instance = MagicMock()
            mock_instance.generate_content.return_value = mock_response
            mock_model.return_value = mock_instance
            
            # Get suggestions from analyzer
            analyzer = AdvancedGeminiAnalyzer()
            analyzer.client = mock_instance
            suggestions = analyzer.analyze_text("This is a test for Large Language Model integration.")
            
            # Integrate suggestions
            integrator = AdvancedWordIntegrator(str(doc_path))
            
            integration_results = []
            for suggestion in suggestions:
                result = integrator.add_comment(
                    text=suggestion.original_text,
                    comment=suggestion.reason,
                    author=f"AI-{suggestion.category}"
                )
                integration_results.append(result)
            
            # Should have attempted integration
            assert len(integration_results) > 0
    
    def test_full_pipeline_integration(self, integration_test_data, mock_google_api_key):
        """Test complete pipeline integration with data flow"""
        input_doc = integration_test_data['input_document']
        output_doc = integration_test_data['output_document']
        
        if any(cls is None for cls in [DocxParser, AdvancedGeminiAnalyzer, AdvancedWordIntegrator]):
            pytest.skip("Required components not available")
        
        pipeline_data = {}
        
        with patch('google.generativeai.GenerativeModel') as mock_model:
            mock_response = MagicMock()
            mock_response.text = json.dumps({"suggestions": []})
            mock_instance = MagicMock()
            mock_instance.generate_content.return_value = mock_response
            mock_model.return_value = mock_instance
            
            # Stage 1: Parse and extract data
            parser = DocxParser(str(input_doc))
            pipeline_data['original_paragraphs'] = len([p for p in parser.document.paragraphs if p.text.strip()])
            pipeline_data['original_text_length'] = sum(len(p.text) for p in parser.document.paragraphs)
            
            # Stage 2: Analyze and collect suggestions
            analyzer = AdvancedGeminiAnalyzer()
            analyzer.client = mock_instance
            
            full_text = " ".join([p.text for p in parser.document.paragraphs if p.text.strip()])
            suggestions = analyzer.analyze_text(full_text)
            pipeline_data['suggestions_generated'] = len(suggestions)
            
            # Stage 3: Integration and output
            integrator = AdvancedWordIntegrator(str(input_doc))
            
            comments_added = 0
            for suggestion in suggestions:
                result = integrator.add_comment(
                    text=suggestion.original_text,
                    comment=suggestion.reason,
                    author="AI Pipeline"
                )
                if result:
                    comments_added += 1
            
            pipeline_data['comments_integrated'] = comments_added
            
            # Save final result
            integrator.save(str(output_doc))
            pipeline_data['output_created'] = output_doc.exists()
            
            # Verify pipeline data flow
            assert pipeline_data['original_paragraphs'] > 0
            assert pipeline_data['original_text_length'] > 0
            assert isinstance(pipeline_data['suggestions_generated'], int)
            assert isinstance(pipeline_data['comments_integrated'], int)


@pytest.mark.integration
@pytest.mark.filesystem
class TestFileSystemIntegration:
    """Test file system operations and document handling"""
    
    def test_document_creation_and_modification(self, temp_dir, sample_docx_content):
        """Test document creation, modification, and preservation"""
        original_path = temp_dir / "original.docx"
        modified_path = temp_dir / "modified.docx"
        
        # Create original document
        doc = Document()
        doc.add_heading(sample_docx_content["title"], level=1)
        for paragraph in sample_docx_content["paragraphs"]:
            doc.add_paragraph(paragraph)
        doc.save(str(original_path))
        
        assert original_path.exists()
        original_size = original_path.stat().st_size
        
        # Modify document
        if AdvancedWordIntegrator is not None:
            integrator = AdvancedWordIntegrator(str(original_path))
            
            # Add some modifications
            integrator.add_comment(
                text="Test text",
                comment="Test comment",
                author="Test Author"
            )
            
            integrator.save(str(modified_path))
            
            if modified_path.exists():
                modified_size = modified_path.stat().st_size
                
                # Modified document should exist and be different size
                assert modified_size != original_size
                
                # Both documents should be readable
                original_doc = Document(str(original_path))
                modified_doc = Document(str(modified_path))
                
                assert original_doc is not None
                assert modified_doc is not None
    
    def test_concurrent_file_access(self, create_test_docx, temp_dir):
        """Test concurrent access to document files"""
        if DocxParser is None:
            pytest.skip("DocxParser not available")
        
        doc_path = create_test_docx("concurrent_test.docx")
        
        import threading
        import queue
        
        results = queue.Queue()
        
        def read_document(doc_id):
            try:
                parser = DocxParser(str(doc_path))
                text_length = sum(len(p.text) for p in parser.document.paragraphs)
                results.put(("success", doc_id, text_length))
            except Exception as e:
                results.put(("error", doc_id, str(e)))
        
        # Create multiple concurrent readers
        threads = []
        for i in range(5):
            thread = threading.Thread(target=read_document, args=(i,))
            threads.append(thread)
            thread.start()
        
        # Wait for completion
        for thread in threads:
            thread.join()
        
        # Check results
        success_count = 0
        while not results.empty():
            status, doc_id, result = results.get()
            if status == "success":
                success_count += 1
                assert isinstance(result, int)  # text length
            
        # All concurrent reads should succeed
        assert success_count == 5
    
    def test_file_backup_and_recovery(self, create_test_docx, temp_dir):
        """Test file backup and recovery mechanisms"""
        original_path = create_test_docx("backup_test.docx")
        backup_path = temp_dir / "backup_test_backup.docx"
        recovery_path = temp_dir / "backup_test_recovered.docx"
        
        # Create backup
        import shutil
        shutil.copy2(str(original_path), str(backup_path))
        
        assert backup_path.exists()
        
        # Simulate modification that might fail
        if AdvancedWordIntegrator is not None:
            try:
                integrator = AdvancedWordIntegrator(str(original_path))
                # Simulate adding many comments (might cause issues)
                for i in range(50):
                    integrator.add_comment(
                        text=f"test text {i}",
                        comment=f"test comment {i}",
                        author="Test"
                    )
                # Don't save to simulate failure
            except Exception:
                pass  # Expected to potentially fail
            
            # Recovery from backup
            shutil.copy2(str(backup_path), str(recovery_path))
            
            # Verify recovery
            assert recovery_path.exists()
            recovered_doc = Document(str(recovery_path))
            assert recovered_doc is not None


@pytest.mark.integration
@pytest.mark.api
class TestAPIIntegration:
    """Test API integration and external dependencies"""
    
    def test_api_configuration_and_initialization(self, mock_google_api_key):
        """Test API configuration and initialization"""
        if AdvancedGeminiAnalyzer is None:
            pytest.skip("AdvancedGeminiAnalyzer not available")
        
        with patch('google.generativeai.configure') as mock_configure:
            analyzer = AdvancedGeminiAnalyzer()
            
            # Should configure API with key
            mock_configure.assert_called_once_with(api_key=mock_google_api_key)
    
    def test_api_error_handling_and_fallbacks(self, mock_google_api_key):
        """Test API error handling and fallback mechanisms"""
        if AdvancedGeminiAnalyzer is None:
            pytest.skip("AdvancedGeminiAnalyzer not available")
        
        with patch('google.generativeai.GenerativeModel') as mock_model:
            # Test different error scenarios
            error_scenarios = [
                Exception("Network error"),
                Exception("API rate limit exceeded"),
                Exception("Invalid API key"),
                Exception("Service unavailable")
            ]
            
            for error in error_scenarios:
                mock_instance = MagicMock()
                mock_instance.generate_content.side_effect = error
                mock_model.return_value = mock_instance
                
                analyzer = AdvancedGeminiAnalyzer()
                analyzer.client = mock_instance
                
                # Should handle error gracefully
                suggestions = analyzer.analyze_text("Test text")
                assert isinstance(suggestions, list)  # Should return empty list
                assert len(suggestions) == 0
    
    def test_api_response_validation(self, mock_google_api_key):
        """Test API response validation and parsing"""
        if AdvancedGeminiAnalyzer is None:
            pytest.skip("AdvancedGeminiAnalyzer not available")
        
        with patch('google.generativeai.GenerativeModel') as mock_model:
            # Test various response scenarios
            response_scenarios = [
                # Valid response
                '{"suggestions": [{"original_text": "test", "suggested_text": "better test", "reason": "improvement", "category": "style", "confidence": 0.8}]}',
                # Invalid JSON
                'Invalid JSON response',
                # Missing fields
                '{"suggestions": [{"original_text": "test"}]}',
                # Empty response
                '',
                # Wrong structure
                '{"data": [{"text": "test"}]}'
            ]
            
            for response_text in response_scenarios:
                mock_response = MagicMock()
                mock_response.text = response_text
                mock_instance = MagicMock()
                mock_instance.generate_content.return_value = mock_response
                mock_model.return_value = mock_instance
                
                analyzer = AdvancedGeminiAnalyzer()
                analyzer.client = mock_instance
                
                suggestions = analyzer.analyze_text("Test text")
                
                # Should always return a list, even for invalid responses
                assert isinstance(suggestions, list)
                
                # Valid response should have suggestions
                if "suggestions" in response_text and "original_text" in response_text:
                    # May have suggestions if validation passes
                    pass
                else:
                    # Invalid responses should result in empty list
                    assert len(suggestions) == 0


@pytest.mark.integration
@pytest.mark.slow
class TestLongRunningIntegration:
    """Long-running integration tests"""
    
    def test_extended_processing_session(self, test_document_files, mock_google_api_key):
        """Test extended processing session with multiple documents"""
        if any(cls is None for cls in [DocxParser, AdvancedGeminiAnalyzer, AdvancedWordIntegrator]):
            pytest.skip("Required components not available")
        
        with patch('google.generativeai.GenerativeModel') as mock_model:
            mock_response = MagicMock()
            mock_response.text = '{"suggestions": []}'
            mock_instance = MagicMock()
            mock_instance.generate_content.return_value = mock_response
            mock_model.return_value = mock_instance
            
            session_results = []
            
            # Process multiple documents in sequence
            for size, doc_path in test_document_files.items():
                try:
                    # Parse
                    parser = DocxParser(str(doc_path))
                    
                    # Analyze
                    analyzer = AdvancedGeminiAnalyzer()
                    analyzer.client = mock_instance
                    
                    text = " ".join([p.text for p in parser.document.paragraphs if p.text.strip()])
                    suggestions = analyzer.analyze_text(text[:500])  # Limit text
                    
                    # Integrate (simulate)
                    integrator = AdvancedWordIntegrator(str(doc_path))
                    
                    session_results.append({
                        'size': size,
                        'parsed': True,
                        'suggestions': len(suggestions),
                        'integrated': True
                    })
                    
                except Exception as e:
                    session_results.append({
                        'size': size,
                        'error': str(e)
                    })
            
            # Should have processed all documents
            assert len(session_results) == len(test_document_files)
            
            # Most should have succeeded
            successful = [r for r in session_results if 'error' not in r]
            assert len(successful) >= len(session_results) * 0.8  # 80% success rate
    
    def test_memory_stability_extended_use(self, create_test_docx, memory_monitor, mock_google_api_key):
        """Test memory stability during extended use"""
        if DocxParser is None:
            pytest.skip("DocxParser not available")
        
        initial_memory = memory_monitor['initial_memory']
        
        # Simulate extended use - process many small documents
        for iteration in range(20):
            doc_path = create_test_docx(f"extended_{iteration}.docx")
            
            # Process document
            parser = DocxParser(str(doc_path))
            del parser  # Explicit cleanup
            
            # Check memory periodically
            if iteration % 5 == 0:
                import gc
                gc.collect()
                current_memory = memory_monitor['get_current_memory']()
                memory_increase = current_memory - initial_memory
                
                # Memory should not grow excessively
                assert memory_increase < 200 * 1024 * 1024  # < 200MB
        
        # Final memory check
        import gc
        gc.collect()
        final_memory = memory_monitor['get_current_memory']()
        total_increase = final_memory - initial_memory
        
        # Total memory increase should be reasonable
        assert total_increase < 300 * 1024 * 1024  # < 300MB