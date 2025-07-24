"""
Test Utilities and Helper Functions
Common utilities for testing the Bachelor Thesis Correction Tool
"""

import json
import time
import psutil
import tempfile
import shutil
from pathlib import Path
from typing import Dict, List, Any, Optional, Callable, Union, Tuple
from contextlib import contextmanager
from unittest.mock import Mock, MagicMock, patch
from docx import Document
import pytest


class TestTimer:
    """Utility for timing test operations"""
    
    def __init__(self):
        self.times = {}
        self.start_times = {}
    
    def start(self, name: str):
        """Start timing an operation"""
        self.start_times[name] = time.time()
    
    def end(self, name: str) -> float:
        """End timing and return elapsed time"""
        if name not in self.start_times:
            raise ValueError(f"Timer '{name}' was not started")
        
        elapsed = time.time() - self.start_times[name]
        self.times[name] = elapsed
        del self.start_times[name]
        return elapsed
    
    def get_time(self, name: str) -> Optional[float]:
        """Get recorded time for operation"""
        return self.times.get(name)
    
    def get_all_times(self) -> Dict[str, float]:
        """Get all recorded times"""
        return self.times.copy()
    
    @contextmanager
    def time_operation(self, name: str):
        """Context manager for timing operations"""
        self.start(name)
        try:
            yield
        finally:
            self.end(name)


class MemoryMonitor:
    """Utility for monitoring memory usage during tests"""
    
    def __init__(self):
        self.process = psutil.Process()
        self.initial_memory = self.process.memory_info().rss
        self.snapshots = {}
    
    def get_current_memory(self) -> int:
        """Get current memory usage in bytes"""
        return self.process.memory_info().rss
    
    def get_memory_increase(self) -> int:
        """Get memory increase from initial"""
        return self.get_current_memory() - self.initial_memory
    
    def take_snapshot(self, name: str):
        """Take a memory snapshot"""
        self.snapshots[name] = self.get_current_memory()
    
    def get_increase_since_snapshot(self, name: str) -> int:
        """Get memory increase since named snapshot"""
        if name not in self.snapshots:
            raise ValueError(f"Snapshot '{name}' not found")
        return self.get_current_memory() - self.snapshots[name]
    
    def force_gc(self):
        """Force garbage collection"""
        import gc
        gc.collect()
    
    @contextmanager
    def monitor_operation(self, name: str, max_increase_mb: int = 100):
        """Monitor memory usage during operation"""
        initial = self.get_current_memory()
        try:
            yield
        finally:
            final = self.get_current_memory()
            increase = final - initial
            max_increase_bytes = max_increase_mb * 1024 * 1024
            
            if increase > max_increase_bytes:
                pytest.fail(f"Memory increase for '{name}': {increase / 1024 / 1024:.1f}MB "
                           f"exceeds limit of {max_increase_mb}MB")


class MockAPIHelper:
    """Helper for creating consistent API mocks"""
    
    @staticmethod
    def create_gemini_mock(suggestions: List[Dict[str, Any]] = None) -> Mock:
        """Create a mock Gemini API client"""
        if suggestions is None:
            suggestions = [
                {
                    "original_text": "test text",
                    "suggested_text": "improved test text",
                    "reason": "Better clarity",
                    "category": "style",
                    "confidence": 0.8
                }
            ]
        
        mock_client = Mock()
        mock_response = Mock()
        mock_response.text = json.dumps({"suggestions": suggestions})
        mock_client.generate_content.return_value = mock_response
        
        return mock_client
    
    @staticmethod
    def create_failing_gemini_mock(error_message: str = "API Error") -> Mock:
        """Create a mock that fails with specified error"""
        mock_client = Mock()
        mock_client.generate_content.side_effect = Exception(error_message)
        return mock_client
    
    @staticmethod
    def create_invalid_response_mock(invalid_response: str = "Invalid JSON") -> Mock:
        """Create a mock that returns invalid responses"""
        mock_client = Mock()
        mock_response = Mock()
        mock_response.text = invalid_response
        mock_client.generate_content.return_value = mock_response
        return mock_client
    
    @contextmanager
    def patch_gemini_api(self, mock_client: Mock):
        """Context manager to patch the Gemini API"""
        with patch('google.generativeai.GenerativeModel') as mock_model:
            mock_model.return_value = mock_client
            with patch('google.generativeai.configure'):
                yield mock_client


class DocumentTestHelper:
    """Helper for creating and manipulating test documents"""
    
    @staticmethod
    def create_simple_document(title: str, paragraphs: List[str], output_path: Path) -> Path:
        """Create a simple document with title and paragraphs"""
        doc = Document()
        doc.add_heading(title, level=1)
        
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
        
        doc.save(str(output_path))
        return output_path
    
    @staticmethod
    def create_complex_document(output_path: Path) -> Path:
        """Create a complex document with various elements"""
        doc = Document()
        
        # Title
        doc.add_heading('Complex Test Document', level=1)
        
        # Abstract
        doc.add_heading('Abstract', level=2)
        doc.add_paragraph('This document contains various elements for comprehensive testing.')
        
        # Introduction with bullet points
        doc.add_heading('Introduction', level=2)
        doc.add_paragraph('This introduction covers multiple aspects:')
        doc.add_paragraph('• First important point', style='List Bullet')
        doc.add_paragraph('• Second crucial aspect', style='List Bullet')
        doc.add_paragraph('• Third essential element', style='List Bullet')
        
        # Main content with subsections
        doc.add_heading('Main Content', level=2)
        doc.add_heading('Subsection A', level=3)
        doc.add_paragraph('Content for subsection A with technical terms like Machine Learning and AI.')
        
        doc.add_heading('Subsection B', level=3)
        doc.add_paragraph('Content for subsection B discussing Large Language Models and their applications.')
        
        # Conclusion
        doc.add_heading('Conclusion', level=2)
        doc.add_paragraph('This document provides comprehensive test scenarios for various components.')
        
        doc.save(str(output_path))
        return output_path
    
    @staticmethod
    def extract_text_with_positions(doc_path: Path) -> List[Tuple[str, int, int]]:
        """Extract text with position information from document"""
        doc = Document(str(doc_path))
        text_segments = []
        position = 0
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text = paragraph.text
                start_pos = position
                end_pos = position + len(text)
                text_segments.append((text, start_pos, end_pos))
                position = end_pos + 1  # +1 for paragraph separator
        
        return text_segments
    
    @staticmethod
    def count_elements(doc_path: Path) -> Dict[str, int]:
        """Count various document elements"""
        doc = Document(str(doc_path))
        
        counts = {
            'paragraphs': 0,
            'headings': 0,
            'total_words': 0,
            'total_chars': 0
        }
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                counts['paragraphs'] += 1
                if paragraph.style.name.startswith('Heading'):
                    counts['headings'] += 1
                
                words = paragraph.text.split()
                counts['total_words'] += len(words)
                counts['total_chars'] += len(paragraph.text)
        
        return counts


class TestDataValidator:
    """Utility for validating test data and results"""
    
    @staticmethod
    def validate_suggestion(suggestion: Any) -> bool:
        """Validate that an object has suggestion-like properties"""
        required_attrs = ['original_text', 'suggested_text', 'reason', 'category', 'confidence']
        
        for attr in required_attrs:
            if not hasattr(suggestion, attr):
                return False
            
            value = getattr(suggestion, attr)
            if value is None or (isinstance(value, str) and not value.strip()):
                return False
        
        # Validate confidence is between 0 and 1
        if not (0 <= suggestion.confidence <= 1):
            return False
        
        return True
    
    @staticmethod
    def validate_suggestion_list(suggestions: List[Any]) -> Dict[str, Any]:
        """Validate a list of suggestions and return statistics"""
        if not isinstance(suggestions, list):
            return {'valid': False, 'error': 'Not a list'}
        
        stats = {
            'valid': True,
            'total_count': len(suggestions),
            'valid_count': 0,
            'invalid_count': 0,
            'categories': {},
            'confidence_stats': {'min': 1.0, 'max': 0.0, 'avg': 0.0},
            'errors': []
        }
        
        confidence_values = []
        
        for i, suggestion in enumerate(suggestions):
            if TestDataValidator.validate_suggestion(suggestion):
                stats['valid_count'] += 1
                
                # Category statistics
                category = suggestion.category
                stats['categories'][category] = stats['categories'].get(category, 0) + 1
                
                # Confidence statistics
                confidence_values.append(suggestion.confidence)
            else:
                stats['invalid_count'] += 1
                stats['errors'].append(f"Invalid suggestion at index {i}")
        
        if confidence_values:
            stats['confidence_stats']['min'] = min(confidence_values)
            stats['confidence_stats']['max'] = max(confidence_values)
            stats['confidence_stats']['avg'] = sum(confidence_values) / len(confidence_values)
        
        if stats['invalid_count'] > 0:
            stats['valid'] = False
        
        return stats
    
    @staticmethod
    def validate_document_chunk(chunk: Any) -> bool:
        """Validate that an object has document chunk properties"""
        required_attrs = ['text', 'start_pos', 'end_pos', 'paragraph_idx']
        
        for attr in required_attrs:
            if not hasattr(chunk, attr):
                return False
        
        # Validate positions
        if chunk.end_pos <= chunk.start_pos:
            return False
        
        if chunk.paragraph_idx < 0:
            return False
        
        return True


class TestEnvironmentHelper:
    """Helper for managing test environment"""
    
    @staticmethod
    @contextmanager
    def temporary_environment_vars(**env_vars):
        """Temporarily set environment variables"""
        import os
        original_values = {}
        
        for key, value in env_vars.items():
            original_values[key] = os.environ.get(key)
            os.environ[key] = value
        
        try:
            yield
        finally:
            for key, original_value in original_values.items():
                if original_value is None:
                    os.environ.pop(key, None)
                else:
                    os.environ[key] = original_value
    
    @staticmethod
    @contextmanager
    def temporary_directory():
        """Create and cleanup temporary directory"""
        temp_dir = Path(tempfile.mkdtemp())
        try:
            yield temp_dir
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)
    
    @staticmethod
    def skip_if_no_api_key():
        """Skip test if API key is not available"""
        import os
        if not os.getenv('GOOGLE_API_KEY'):
            pytest.skip("GOOGLE_API_KEY not available")
    
    @staticmethod
    def skip_if_modules_missing(*module_names):
        """Skip test if required modules are not available"""
        for module_name in module_names:
            try:
                __import__(module_name)
            except ImportError:
                pytest.skip(f"Required module '{module_name}' not available")


class PerformanceTestHelper:
    """Helper utilities for performance testing"""
    
    @staticmethod
    def benchmark_function(func: Callable, *args, iterations: int = 10, **kwargs) -> Dict[str, Any]:
        """Benchmark a function call"""
        times = []
        
        for _ in range(iterations):
            start_time = time.time()
            result = func(*args, **kwargs)
            end_time = time.time()
            times.append(end_time - start_time)
        
        return {
            'times': times,
            'min_time': min(times),
            'max_time': max(times),
            'avg_time': sum(times) / len(times),
            'total_time': sum(times),
            'iterations': iterations,
            'last_result': result
        }
    
    @staticmethod
    def assert_performance_limit(elapsed_time: float, max_time: float, operation_name: str = "Operation"):
        """Assert that operation completed within time limit"""
        if elapsed_time > max_time:
            pytest.fail(f"{operation_name} took {elapsed_time:.2f}s, "
                       f"exceeding limit of {max_time:.2f}s")
    
    @staticmethod
    def assert_memory_limit(memory_usage: int, max_memory_mb: int, operation_name: str = "Operation"):
        """Assert that operation used memory within limit"""
        max_memory_bytes = max_memory_mb * 1024 * 1024
        if memory_usage > max_memory_bytes:
            used_mb = memory_usage / 1024 / 1024
            pytest.fail(f"{operation_name} used {used_mb:.1f}MB, "
                       f"exceeding limit of {max_memory_mb}MB")


class MockDataGenerator:
    """Generate mock data for testing"""
    
    @staticmethod
    def generate_text_samples(count: int, min_length: int = 50, max_length: int = 200) -> List[str]:
        """Generate sample text strings"""
        import random
        
        base_words = [
            'Forschung', 'Technologie', 'Analyse', 'Methodik', 'Ergebnis', 'Diskussion',
            'Künstliche', 'Intelligenz', 'Machine', 'Learning', 'Algorithmus', 'Daten',
            'Wissenschaft', 'Entwicklung', 'Implementation', 'Evaluation', 'System'
        ]
        
        samples = []
        for _ in range(count):
            word_count = random.randint(min_length // 10, max_length // 10)
            words = [random.choice(base_words) for _ in range(word_count)]
            text = ' '.join(words) + '.'
            samples.append(text)
        
        return samples
    
    @staticmethod
    def generate_api_responses(count: int, categories: Optional[List[str]] = None) -> List[Dict]:
        """Generate mock API response data"""
        import random
        
        if categories is None:
            categories = ['grammar', 'style', 'clarity', 'academic']
        
        responses = []
        for _ in range(count):
            suggestions = []
            for _ in range(random.randint(1, 5)):
                suggestion = {
                    "original_text": f"sample text {random.randint(1, 100)}",
                    "suggested_text": f"improved text {random.randint(1, 100)}",
                    "reason": f"test reason {random.randint(1, 50)}",
                    "category": random.choice(categories),
                    "confidence": round(random.uniform(0.7, 0.95), 2)
                }
                suggestions.append(suggestion)
            
            responses.append({"suggestions": suggestions})
        
        return responses


class TestReportHelper:
    """Helper for generating test reports and summaries"""
    
    @staticmethod
    def generate_test_summary(test_results: Dict[str, Any]) -> str:
        """Generate a summary of test results"""
        summary = []
        summary.append("=" * 50)
        summary.append("TEST EXECUTION SUMMARY")
        summary.append("=" * 50)
        
        if 'timing' in test_results:
            summary.append("\nTiming Results:")
            for operation, time_taken in test_results['timing'].items():
                summary.append(f"  {operation}: {time_taken:.3f}s")
        
        if 'memory' in test_results:
            summary.append("\nMemory Usage:")
            for operation, memory_used in test_results['memory'].items():
                mb_used = memory_used / 1024 / 1024
                summary.append(f"  {operation}: {mb_used:.1f}MB")
        
        if 'validation' in test_results:
            summary.append("\nValidation Results:")
            for check, result in test_results['validation'].items():
                status = "PASS" if result else "FAIL"
                summary.append(f"  {check}: {status}")
        
        return "\n".join(summary)
    
    @staticmethod
    def save_test_artifacts(artifacts: Dict[str, Any], output_dir: Path):
        """Save test artifacts to directory"""
        output_dir.mkdir(parents=True, exist_ok=True)
        
        for name, data in artifacts.items():
            file_path = output_dir / f"{name}.json"
            
            if isinstance(data, (dict, list)):
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2, ensure_ascii=False)
            else:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(str(data))


# Decorator utilities for tests
def requires_api_key(func):
    """Decorator to skip test if API key is not available"""
    def wrapper(*args, **kwargs):
        TestEnvironmentHelper.skip_if_no_api_key()
        return func(*args, **kwargs)
    return wrapper


def requires_modules(*modules):
    """Decorator to skip test if modules are not available"""
    def decorator(func):
        def wrapper(*args, **kwargs):
            TestEnvironmentHelper.skip_if_modules_missing(*modules)
            return func(*args, **kwargs)
        return wrapper
    return decorator


def monitor_performance(max_time: float = None, max_memory_mb: int = None):
    """Decorator to monitor test performance"""
    def decorator(func):
        def wrapper(*args, **kwargs):
            timer = TestTimer()
            memory_monitor = MemoryMonitor()
            
            with timer.time_operation('test_execution'):
                with memory_monitor.monitor_operation('test_execution', max_memory_mb or 100):
                    result = func(*args, **kwargs)
            
            elapsed = timer.get_time('test_execution')
            if max_time and elapsed > max_time:
                pytest.fail(f"Test exceeded time limit: {elapsed:.2f}s > {max_time:.2f}s")
            
            return result
        return wrapper
    return decorator