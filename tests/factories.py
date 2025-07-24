"""
Test Data Factories
Factory patterns for creating test data objects consistently across tests
"""

import random
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document

# Import project modules for type hints
try:
    from src.analyzers.advanced_gemini_analyzer import Suggestion
    from src.parsers.docx_parser import DocumentChunk
except ImportError:
    # Handle missing modules gracefully
    Suggestion = None
    DocumentChunk = None


class SuggestionFactory:
    """Factory for creating Suggestion objects"""
    
    CATEGORIES = ['grammar', 'style', 'clarity', 'academic']
    
    SAMPLE_TEXTS = {
        'grammar': [
            'Dies ist falsch',
            'Das Verb stimmt nicht überein',
            'Komma fehlt hier',
            'Rechtschreibfehler vorhanden'
        ],
        'style': [
            'umgangssprachliche Formulierung',
            'zu informeller Ton',
            'Wiederholung vermeiden',
            'präzisere Wortwahl nötig'
        ],
        'clarity': [
            'unklar formuliert',
            'mehrdeutige Aussage',
            'komplexer Satzbau',
            'Verständlichkeit verbessern'
        ],
        'academic': [
            'Large Language Model',
            'Machine Learning Algorithmus',
            'KI-basierte Technologie',
            'wissenschaftliche Methodik'
        ]
    }
    
    SAMPLE_IMPROVEMENTS = {
        'grammar': [
            'Dies ist korrekt',
            'Das Verb stimmt mit dem Subjekt überein',
            'Komma korrekt gesetzt',
            'Rechtschreibung korrigiert'
        ],
        'style': [
            'wissenschaftliche Formulierung',
            'angemessener akademischer Ton',
            'Variation in der Wortwahl',
            'präzise und treffende Wortwahl'
        ],
        'clarity': [
            'klar und eindeutig formuliert',
            'unmissverständliche Aussage',
            'verständlicher Satzbau',
            'verbesserte Verständlichkeit'
        ],
        'academic': [
            'Large Language Model (LLM)',
            'Machine Learning Algorithmus mit spezifischem Anwendungsbereich',
            'KI-basierte Technologie der nächsten Generation',
            'evidenzbasierte wissenschaftliche Methodik'
        ]
    }
    
    SAMPLE_REASONS = {
        'grammar': [
            'Grammatikfehler korrigieren',
            'Satzstruktur verbessern',
            'Interpunktion anpassen',
            'Rechtschreibung korrigieren'
        ],
        'style': [
            'Stilistische Verbesserung',
            'Angemessenen Ton verwenden',
            'Redundanz vermeiden',
            'Präzision erhöhen'
        ],
        'clarity': [
            'Klarheit verbessern',
            'Verständlichkeit erhöhen',
            'Mehrdeutigkeit beseitigen',
            'Lesbarkeit optimieren'
        ],
        'academic': [
            'Wissenschaftliche Präzision',
            'Fachterminologie erklären',
            'Akademische Standards einhalten',
            'Methodische Genauigkeit'
        ]
    }
    
    @classmethod
    def create(cls, 
               original_text: Optional[str] = None,
               suggested_text: Optional[str] = None,
               reason: Optional[str] = None,
               category: Optional[str] = None,
               confidence: Optional[float] = None,
               position: Optional[Tuple[int, int]] = None) -> 'Suggestion':
        """Create a Suggestion object with provided or random data"""
        
        if Suggestion is None:
            # Return mock object if Suggestion class not available
            @dataclass
            class MockSuggestion:
                original_text: str
                suggested_text: str
                reason: str
                category: str
                confidence: float
                position: Tuple[int, int]
            
            Suggestion = MockSuggestion
        
        # Use provided values or generate random ones
        category = category or random.choice(cls.CATEGORIES)
        original_text = original_text or random.choice(cls.SAMPLE_TEXTS[category])
        suggested_text = suggested_text or random.choice(cls.SAMPLE_IMPROVEMENTS[category])
        reason = reason or random.choice(cls.SAMPLE_REASONS[category])
        confidence = confidence or round(random.uniform(0.7, 0.95), 2)
        position = position or (random.randint(0, 1000), random.randint(1001, 2000))
        
        return Suggestion(
            original_text=original_text,
            suggested_text=suggested_text,
            reason=reason,
            category=category,
            confidence=confidence,
            position=position
        )
    
    @classmethod
    def create_batch(cls, count: int, category: Optional[str] = None) -> List['Suggestion']:
        """Create multiple suggestions"""
        return [cls.create(category=category) for _ in range(count)]
    
    @classmethod
    def create_for_text(cls, text: str, categories: Optional[List[str]] = None) -> List['Suggestion']:
        """Create suggestions that could realistically apply to given text"""
        categories = categories or cls.CATEGORIES
        suggestions = []
        
        words = text.split()
        for i, category in enumerate(categories):
            if i < len(words):
                # Use actual words from the text
                original_word = words[i]
                start_pos = text.find(original_word)
                end_pos = start_pos + len(original_word)
                
                suggestion = cls.create(
                    original_text=original_word,
                    category=category,
                    position=(start_pos, end_pos)
                )
                suggestions.append(suggestion)
        
        return suggestions


class DocumentChunkFactory:
    """Factory for creating DocumentChunk objects"""
    
    SAMPLE_PARAGRAPHS = [
        "Dies ist ein Beispielabsatz für die Dokumentenverarbeitung. Er enthält mehrere Sätze und verschiedene Wörter.",
        "Künstliche Intelligenz und Machine Learning sind wichtige Technologien in der modernen Forschung.",
        "Die Implementierung von Algorithmen erfordert sorgfältige Planung und systematische Herangehensweise.",
        "Wissenschaftliche Arbeiten folgen bestimmten Standards und Konventionen der jeweiligen Disziplin.",
        "Large Language Models haben die Verarbeitung natürlicher Sprache revolutioniert."
    ]
    
    @classmethod
    def create(cls,
               text: Optional[str] = None,
               start_pos: Optional[int] = None,
               end_pos: Optional[int] = None,
               paragraph_idx: Optional[int] = None,
               element_type: str = "paragraph") -> 'DocumentChunk':
        """Create a DocumentChunk with provided or random data"""
        
        if DocumentChunk is None:
            # Return mock object if DocumentChunk class not available
            @dataclass
            class MockDocumentChunk:
                text: str
                start_pos: int
                end_pos: int
                paragraph_idx: int
                element_type: str
                suggestions: List = None
                
                def __post_init__(self):
                    if self.suggestions is None:
                        self.suggestions = []
            
            DocumentChunk = MockDocumentChunk
        
        # Use provided values or generate random ones
        text = text or random.choice(cls.SAMPLE_PARAGRAPHS)
        start_pos = start_pos or random.randint(0, 1000)
        end_pos = end_pos or (start_pos + len(text))
        paragraph_idx = paragraph_idx or random.randint(0, 10)
        
        chunk = DocumentChunk(
            text=text,
            start_pos=start_pos,
            end_pos=end_pos,
            paragraph_idx=paragraph_idx,
            element_type=element_type
        )
        
        return chunk
    
    @classmethod
    def create_sequence(cls, count: int, base_text: Optional[str] = None) -> List['DocumentChunk']:
        """Create a sequence of chunks that could be from the same document"""
        if base_text:
            # Split base text into chunks
            chunk_size = len(base_text) // count
            chunks = []
            for i in range(count):
                start = i * chunk_size
                end = start + chunk_size if i < count - 1 else len(base_text)
                chunk_text = base_text[start:end]
                
                chunk = cls.create(
                    text=chunk_text,
                    start_pos=start,
                    end_pos=end,
                    paragraph_idx=i
                )
                chunks.append(chunk)
            return chunks
        else:
            # Create sequence with continuous positions
            chunks = []
            current_pos = 0
            for i in range(count):
                text = random.choice(cls.SAMPLE_PARAGRAPHS)
                chunk = cls.create(
                    text=text,
                    start_pos=current_pos,
                    end_pos=current_pos + len(text),
                    paragraph_idx=i
                )
                chunks.append(chunk)
                current_pos += len(text) + 2  # Add space between chunks
            return chunks


class DocumentFactory:
    """Factory for creating test documents"""
    
    DOCUMENT_TEMPLATES = {
        'academic_paper': {
            'title': 'Artificial Intelligence in Modern Research',
            'sections': [
                'Abstract',
                'Introduction',
                'Literature Review',
                'Methodology',
                'Results',
                'Discussion',
                'Conclusion',
                'References'
            ],
            'content_per_section': [
                'This paper explores the application of artificial intelligence in contemporary research.',
                'Artificial intelligence has become increasingly important in various research domains.',
                'Previous studies have shown significant advances in machine learning applications.',
                'Our methodology employs state-of-the-art algorithms for data analysis.',
                'The results demonstrate significant improvements over baseline approaches.',
                'These findings have important implications for future research directions.',
                'In conclusion, AI technologies offer promising opportunities for research advancement.',
                'References to relevant literature and sources are provided below.'
            ]
        },
        'thesis_chapter': {
            'title': 'Implementation of Language Models',
            'sections': [
                'Introduction',
                'Background',
                'Implementation Details',
                'Evaluation',
                'Summary'
            ],
            'content_per_section': [
                'This chapter presents the implementation of advanced language models.',
                'Language models have evolved significantly with transformer architectures.',
                'The implementation follows established software engineering principles.',
                'Comprehensive evaluation demonstrates the effectiveness of our approach.',
                'This chapter summarizes the key contributions and findings.'
            ]
        },
        'research_proposal': {
            'title': 'Research Proposal: AI-Enhanced Document Processing',
            'sections': [
                'Problem Statement',
                'Research Questions',
                'Proposed Methodology',
                'Expected Outcomes',
                'Timeline'
            ],
            'content_per_section': [
                'Current document processing systems lack intelligent analysis capabilities.',
                'How can AI improve accuracy and efficiency of document processing?',
                'We propose a multi-stage approach using deep learning techniques.',
                'Expected outcomes include improved processing accuracy and user satisfaction.',
                'The research timeline spans 12 months with quarterly milestones.'
            ]
        }
    }
    
    @classmethod
    def create_docx(cls, 
                   template: str = 'academic_paper',
                   output_path: Optional[Path] = None,
                   add_errors: bool = False) -> Path:
        """Create a DOCX document from template"""
        
        if template not in cls.DOCUMENT_TEMPLATES:
            template = 'academic_paper'
        
        template_data = cls.DOCUMENT_TEMPLATES[template]
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading(template_data['title'], level=1)
        
        # Add sections
        for i, (section, content) in enumerate(zip(template_data['sections'], 
                                                  template_data['content_per_section'])):
            # Add section heading
            doc.add_heading(section, level=2)
            
            # Add content with optional errors
            content_text = content
            if add_errors and random.random() < 0.3:  # 30% chance of error
                content_text = cls._introduce_error(content_text)
            
            doc.add_paragraph(content_text)
            
            # Add additional paragraphs for some sections
            if section in ['Introduction', 'Methodology', 'Results']:
                additional_content = f"Additional details for {section.lower()} are provided here. " \
                                   f"This includes more comprehensive information and analysis."
                if add_errors and random.random() < 0.2:  # 20% chance of error
                    additional_content = cls._introduce_error(additional_content)
                doc.add_paragraph(additional_content)
        
        # Save document
        if output_path is None:
            import tempfile
            output_path = Path(tempfile.mktemp(suffix='.docx'))
        
        doc.save(str(output_path))
        return output_path
    
    @classmethod
    def _introduce_error(cls, text: str) -> str:
        """Introduce random errors into text for testing"""
        error_types = [
            lambda t: t.replace(' ist ', ' sind '),  # Grammar error
            lambda t: t.replace('.', ''),  # Missing punctuation
            lambda t: t.replace('Forschung', 'forschung'),  # Capitalization error
            lambda t: t + ' Das ist redundant.',  # Style issue
            lambda t: t.replace(' und ', ' und und '),  # Repetition
        ]
        
        error_function = random.choice(error_types)
        return error_function(text)
    
    @classmethod
    def create_with_known_issues(cls, issue_types: List[str], output_path: Optional[Path] = None) -> Tuple[Path, List[Dict]]:
        """Create document with specific types of known issues for testing"""
        doc = Document()
        doc.add_heading('Test Document with Known Issues', level=1)
        
        known_issues = []
        
        for issue_type in issue_types:
            if issue_type == 'grammar':
                problematic_text = "Dies sind ein grammatikalischer Fehler in diesem Satz."
                doc.add_paragraph(problematic_text)
                known_issues.append({
                    'type': 'grammar',
                    'text': problematic_text,
                    'issue': 'Subject-verb disagreement',
                    'position': 'paragraph'
                })
            
            elif issue_type == 'style':
                problematic_text = "Das ist mega cool und voll interessant für die Forschung."
                doc.add_paragraph(problematic_text)
                known_issues.append({
                    'type': 'style',
                    'text': problematic_text,
                    'issue': 'Informal language in academic context',
                    'position': 'paragraph'
                })
            
            elif issue_type == 'clarity':
                problematic_text = "Die Sache, die hier gemacht wurde, ist so, dass es funktioniert."
                doc.add_paragraph(problematic_text)
                known_issues.append({
                    'type': 'clarity',
                    'text': problematic_text,
                    'issue': 'Vague and unclear expression',
                    'position': 'paragraph'
                })
            
            elif issue_type == 'academic':
                problematic_text = "KI und ML sind wichtig für NLP Anwendungen."
                doc.add_paragraph(problematic_text)
                known_issues.append({
                    'type': 'academic',
                    'text': problematic_text,
                    'issue': 'Unexplained abbreviations',
                    'position': 'paragraph'
                })
        
        # Save document
        if output_path is None:
            import tempfile
            output_path = Path(tempfile.mktemp(suffix='.docx'))
        
        doc.save(str(output_path))
        return output_path, known_issues


class TestDataFactory:
    """Main factory for creating comprehensive test data sets"""
    
    @classmethod
    def create_analysis_scenario(cls, scenario_type: str = 'standard') -> Dict[str, Any]:
        """Create complete test scenario with document, expected suggestions, etc."""
        
        scenarios = {
            'standard': {
                'description': 'Standard academic document with mixed issues',
                'document_template': 'academic_paper',
                'expected_suggestion_count': (5, 15),
                'expected_categories': ['grammar', 'style', 'clarity', 'academic'],
                'processing_time_limit': 60.0,  # seconds
                'memory_limit': 100 * 1024 * 1024  # 100MB
            },
            'grammar_heavy': {
                'description': 'Document with primarily grammatical issues',
                'document_template': 'thesis_chapter',
                'expected_suggestion_count': (8, 20),
                'expected_categories': ['grammar'],
                'processing_time_limit': 45.0,
                'memory_limit': 50 * 1024 * 1024
            },
            'academic_focus': {
                'description': 'Document requiring academic style improvements',
                'document_template': 'research_proposal',
                'expected_suggestion_count': (3, 10),
                'expected_categories': ['academic', 'style'],
                'processing_time_limit': 30.0,
                'memory_limit': 30 * 1024 * 1024
            },
            'large_document': {
                'description': 'Large document for performance testing',
                'document_template': 'academic_paper',
                'expected_suggestion_count': (20, 50),
                'expected_categories': ['grammar', 'style', 'clarity', 'academic'],
                'processing_time_limit': 180.0,  # 3 minutes
                'memory_limit': 500 * 1024 * 1024  # 500MB
            }
        }
        
        return scenarios.get(scenario_type, scenarios['standard'])
    
    @classmethod
    def create_test_suite_data(cls, temp_dir: Path) -> Dict[str, Any]:
        """Create comprehensive test data for entire test suite"""
        
        test_data = {
            'documents': {},
            'suggestions': {},
            'chunks': {},
            'scenarios': {}
        }
        
        # Create various document types
        for template in DocumentFactory.DOCUMENT_TEMPLATES.keys():
            doc_path = DocumentFactory.create_docx(
                template=template,
                output_path=temp_dir / f"{template}.docx"
            )
            test_data['documents'][template] = doc_path
        
        # Create suggestion sets for different categories
        for category in SuggestionFactory.CATEGORIES:
            test_data['suggestions'][category] = SuggestionFactory.create_batch(5, category)
        
        # Create document chunks
        test_data['chunks']['sequence'] = DocumentChunkFactory.create_sequence(10)
        
        # Create test scenarios
        for scenario_type in ['standard', 'grammar_heavy', 'academic_focus', 'large_document']:
            test_data['scenarios'][scenario_type] = cls.create_analysis_scenario(scenario_type)
        
        return test_data


# Utility functions for test data management
def cleanup_test_data(test_data: Dict[str, Any]):
    """Clean up test data files and objects"""
    if 'documents' in test_data:
        for doc_path in test_data['documents'].values():
            if isinstance(doc_path, Path) and doc_path.exists():
                try:
                    doc_path.unlink()
                except Exception:
                    pass  # Ignore cleanup errors


def validate_test_data(test_data: Dict[str, Any]) -> bool:
    """Validate that test data is properly structured"""
    required_keys = ['documents', 'suggestions', 'chunks', 'scenarios']
    
    for key in required_keys:
        if key not in test_data:
            return False
    
    # Validate documents exist
    if 'documents' in test_data:
        for doc_path in test_data['documents'].values():
            if isinstance(doc_path, Path) and not doc_path.exists():
                return False
    
    return True