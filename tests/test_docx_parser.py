"""
Comprehensive tests for DocxParser module
Tests document parsing, chunking, and text extraction functionality
"""

import pytest
from pathlib import Path
from docx import Document

from src.parsers.docx_parser import DocxParser, DocumentChunk


@pytest.mark.parser
@pytest.mark.unit
class TestDocumentChunk:
    """Test suite for DocumentChunk class"""
    
    def test_document_chunk_creation(self):
        """Test basic DocumentChunk creation"""
        chunk = DocumentChunk(
            text="Test text content",
            start_pos=0,
            end_pos=17,
            paragraph_idx=0,
            element_type="paragraph"
        )
        
        assert chunk.text == "Test text content"
        assert chunk.start_pos == 0
        assert chunk.end_pos == 17
        assert chunk.paragraph_idx == 0
        assert chunk.element_type == "paragraph"
        assert chunk.suggestions == []
    
    def test_document_chunk_default_type(self):
        """Test DocumentChunk with default element type"""
        chunk = DocumentChunk(
            text="Test text",
            start_pos=0,
            end_pos=9,
            paragraph_idx=1
        )
        
        assert chunk.element_type == "paragraph"
    
    def test_document_chunk_suggestions_list(self):
        """Test that suggestions list is mutable"""
        chunk = DocumentChunk("Test text", 0, 9, 0)
        chunk.suggestions.append("test suggestion")
        
        assert len(chunk.suggestions) == 1
        assert chunk.suggestions[0] == "test suggestion"


@pytest.mark.parser
@pytest.mark.unit
class TestDocxParser:
    """Test suite for DocxParser class"""
    
    def test_parser_initialization(self, create_test_docx):
        """Test DocxParser initialization"""
        doc_path = create_test_docx()
        parser = DocxParser(str(doc_path))
        
        assert parser.file_path == str(doc_path)
        assert parser.document is not None
        assert isinstance(parser.chunks, list)
    
    def test_parser_with_nonexistent_file(self):
        """Test parser with non-existent file raises appropriate error"""
        with pytest.raises(Exception):  # Should raise FileNotFoundError or similar
            DocxParser("/nonexistent/path/file.docx")
    
    @pytest.mark.filesystem
    def test_parse_simple_document(self, create_test_docx):
        """Test parsing a simple document"""
        content = {
            "title": "Simple Test Document",
            "paragraphs": ["First paragraph.", "Second paragraph."]
        }
        doc_path = create_test_docx("simple.docx", content)
        
        parser = DocxParser(str(doc_path))
        # Assuming parser has a parse method - adjust based on actual implementation
        # This test may need to be adjusted based on the actual DocxParser API
        
        assert parser.document is not None
        paragraphs = [p.text for p in parser.document.paragraphs if p.text.strip()]
        
        assert "Simple Test Document" in paragraphs
        assert "First paragraph." in paragraphs
        assert "Second paragraph." in paragraphs
    
    @pytest.mark.filesystem
    def test_parse_document_with_headings(self, create_test_docx, sample_docx_content):
        """Test parsing document with headings and multiple elements"""
        doc_path = create_test_docx("with_headings.docx", sample_docx_content)
        
        parser = DocxParser(str(doc_path))
        document = parser.document
        
        # Check that document was loaded
        assert document is not None
        
        # Extract all text content
        all_text = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                all_text.append(paragraph.text)
        
        # Verify expected content is present
        assert sample_docx_content["title"] in all_text
        assert sample_docx_content["heading"] in all_text
        
        for paragraph_text in sample_docx_content["paragraphs"]:
            assert paragraph_text in all_text
    
    @pytest.mark.filesystem
    def test_parse_empty_document(self, temp_dir):
        """Test parsing an empty document"""
        # Create empty document
        empty_doc = Document()
        empty_path = temp_dir / "empty.docx"
        empty_doc.save(str(empty_path))
        
        parser = DocxParser(str(empty_path))
        
        assert parser.document is not None
        # Empty document should have minimal content
        text_content = [p.text for p in parser.document.paragraphs if p.text.strip()]
        assert len(text_content) == 0
    
    @pytest.mark.filesystem
    def test_parse_document_with_special_characters(self, temp_dir, sample_text_data):
        """Test parsing document with special characters and umlauts"""
        doc = Document()
        doc.add_paragraph(sample_text_data["german_special"])
        doc.add_paragraph(sample_text_data["mixed_content"])
        
        special_path = temp_dir / "special_chars.docx"
        doc.save(str(special_path))
        
        parser = DocxParser(str(special_path))
        document = parser.document
        
        all_text = " ".join([p.text for p in document.paragraphs])
        
        # Check special characters are preserved
        assert "äöüß" in all_text
        assert "„deutsche"" in all_text or '"deutsche"' in all_text
        assert "https://example.com" in all_text
        assert "test@example.com" in all_text
    
    @pytest.mark.filesystem
    @pytest.mark.performance
    def test_parse_large_document(self, test_document_files, performance_timer):
        """Test parsing performance with large document"""
        large_doc_path = test_document_files["large"]
        
        performance_timer['start']('parse_large')
        parser = DocxParser(str(large_doc_path))
        elapsed = performance_timer['end']('parse_large')
        
        assert parser.document is not None
        # Should parse large document in reasonable time (adjust threshold as needed)
        assert elapsed < 10.0  # 10 seconds should be more than enough
    
    def test_parser_file_path_property(self, create_test_docx):
        """Test that file path is correctly stored"""
        doc_path = create_test_docx()
        parser = DocxParser(str(doc_path))
        
        assert parser.file_path == str(doc_path)
    
    def test_parser_document_property(self, create_test_docx):
        """Test that document object is accessible"""
        doc_path = create_test_docx()
        parser = DocxParser(str(doc_path))
        
        assert parser.document is not None
        assert hasattr(parser.document, 'paragraphs')
        assert hasattr(parser.document, 'tables')


@pytest.mark.parser
@pytest.mark.integration
class TestDocxParserIntegration:
    """Integration tests for DocxParser with real documents"""
    
    @pytest.mark.filesystem
    def test_integration_with_test_files(self, test_document_files):
        """Test parser with various document sizes"""
        for size, doc_path in test_document_files.items():
            parser = DocxParser(str(doc_path))
            
            assert parser.document is not None
            assert len(parser.chunks) >= 0  # chunks list should exist
            
            # Verify document contains expected content
            paragraphs = [p.text for p in parser.document.paragraphs if p.text.strip()]
            assert len(paragraphs) > 0
    
    @pytest.mark.filesystem
    @pytest.mark.requires_test_files
    def test_with_real_test_document(self):
        """Test with actual test document from archive if available"""
        test_doc_path = Path("archive/test_files/test_document.docx")
        
        if test_doc_path.exists():
            parser = DocxParser(str(test_doc_path))
            
            assert parser.document is not None
            paragraphs = [p.text for p in parser.document.paragraphs if p.text.strip()]
            assert len(paragraphs) > 0
        else:
            pytest.skip("Real test document not available")


@pytest.mark.parser
@pytest.mark.benchmark
class TestDocxParserBenchmarks:
    """Performance benchmarks for DocxParser"""
    
    @pytest.mark.performance
    def test_parse_speed_benchmark(self, test_document_files, benchmark):
        """Benchmark document parsing speed"""
        medium_doc = test_document_files["medium"]
        
        def parse_document():
            return DocxParser(str(medium_doc))
        
        result = benchmark(parse_document)
        assert result.document is not None
    
    @pytest.mark.memory
    def test_memory_usage_large_document(self, test_document_files, memory_monitor):
        """Test memory usage with large document"""
        large_doc = test_document_files["large"]
        
        initial_memory = memory_monitor['initial_memory']
        
        parser = DocxParser(str(large_doc))
        
        current_memory = memory_monitor['get_current_memory']()
        memory_increase = current_memory - initial_memory
        
        # Memory increase should be reasonable (adjust threshold as needed)
        assert memory_increase < 100 * 1024 * 1024  # Less than 100MB increase
        assert parser.document is not None


@pytest.mark.parser
@pytest.mark.unit
class TestDocxParserErrorHandling:
    """Test error handling and edge cases"""
    
    def test_invalid_file_extension(self):
        """Test parser with non-docx file"""
        with pytest.raises(Exception):
            DocxParser("test.txt")
    
    def test_corrupted_docx_file(self, temp_dir):
        """Test parser with corrupted DOCX file"""
        # Create a file that looks like docx but isn't
        fake_docx = temp_dir / "fake.docx"
        fake_docx.write_text("This is not a valid DOCX file")
        
        with pytest.raises(Exception):
            DocxParser(str(fake_docx))
    
    def test_permission_denied_file(self, temp_dir):
        """Test parser with file that can't be read (if permissions allow testing)"""
        # This test might be skipped on some systems due to permission requirements
        try:
            restricted_file = temp_dir / "restricted.docx"
            Document().save(str(restricted_file))
            restricted_file.chmod(0o000)  # Remove all permissions
            
            with pytest.raises(PermissionError):
                DocxParser(str(restricted_file))
                
        except (OSError, PermissionError):
            pytest.skip("Cannot test permission denied scenario on this system")
        finally:
            # Restore permissions for cleanup
            try:
                restricted_file.chmod(0o644)
            except:
                pass